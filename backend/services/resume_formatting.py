"""Shared resume-formatting primitives: redaction, name masking, and a
compositional PDF/DOCX renderer used by both the standalone Resume
Generator (backend/routers/resume_generator.py) and the older, still-live
KAE-submission fixed-format renderers (backend/routers/kae_submission.py).

One engine, not two — kae_submission.py's 6 named formats are just
compositional configs that call render_resume_pdf() under the hood
(see kae_submission.py's _config_for_style()), so there's a single place
that knows how to build a resume document.
"""
import io
import os
import re
from typing import Optional
from xml.sax.saxutils import escape as _esc


_SUBHEADING_ALLCAPS_RE = re.compile(r'^[A-Z0-9 &/,.\'\-]{3,55}:?$')

# Local to this bolding feature only -- deliberately NOT merged into
# improved_parser.SECTION_HEADERS, which several other parsing functions
# (section-boundary detection, name extraction) depend on for different,
# more consequential behavior. Common mixed-case resume sub-headings that
# the strict ALL-CAPS heuristic below doesn't catch (e.g. "Functional
# Skills", "Competencies:") -- found via a real resume that used exactly
# these as bare one-line headers.
_EXTRA_SUBHEADINGS = frozenset([
    'competencies', 'functional skills', 'system configuration',
    'testing and validation', 'domain experience', 'system cutover',
    'system cutover (go-live)',
])


def _is_subheading(line: str) -> bool:
    """Real improvement (2026-08-18): the body text was rendered as one
    visually flat block, losing all the structure a real resume has --
    bold section headers, bold "Employer:"/"Client:" lines -- confirmed
    directly against a real generated PDF next to its own source document.
    Flags a line for bold/emphasized rendering instead of plain body text.
    Deliberately conservative: only lines matching this codebase's own
    curated SECTION_HEADERS list (plus a small local supplement above), a
    real "Employer:"/"Client:" field label, or a short genuinely ALL-CAPS
    line (a common resume convention for section headers, e.g. "DOMAIN
    EXPERIENCE:") -- never a guess at arbitrary emphasis within ordinary
    sentences."""
    from services.improved_parser import SECTION_HEADERS
    s = line.strip()
    if not s:
        return False
    normalized = s.rstrip(':').strip().lower()
    if normalized in SECTION_HEADERS or normalized in _EXTRA_SUBHEADINGS:
        return True
    low = s.lower()
    if low.startswith('employer:') or low.startswith('client:'):
        return True
    # Real bug fix (2026-08-18): a bare single-word ALL-CAPS fragment like
    # "SAP." was matching this branch -- confirmed directly against a real
    # generated PDF where the tail end of a hard-wrapped sentence ("...
    # advanced migration of\nSAP.") rendered "SAP." as its own bold,
    # spaced-out subheading line instead of the last two words of the
    # sentence it belongs to. Real section headings in this codebase's own
    # convention are always multi-word ("KEY SKILLS", "DOMAIN EXPERIENCE")
    # -- single ALL-CAPS tokens standing alone are acronyms/fragments, not
    # headings, so this branch now requires at least two words.
    if (_SUBHEADING_ALLCAPS_RE.match(s) and any(c.isalpha() for c in s)
            and s.upper() == s and len(s.rstrip(':').split()) >= 2):
        return True
    return False


def mask_name(full_name: str) -> str:
    """'Mohsin Khan' -> 'Mohsin K' (first name + first letter of last name,
    no period — matches the exact rule requested: "Rahul Sharma -> Rahul S").
    Single-word names pass through unchanged; there's no surname to mask."""
    parts = (full_name or "").strip().split()
    if not parts:
        return "Candidate"
    if len(parts) == 1:
        return parts[0]
    return f"{parts[0]} {parts[-1][0]}"


def redact_contact(text: str, candidate: dict, redact_phone: bool, redact_email: bool) -> str:
    """Strip phone/email from free text — conditionally, since a caller that
    IS showing contact details (show_mobile/show_email true) wants them left
    alone in the body, not stripped just because the field exists."""
    if not text:
        return text
    out = text
    phone = (candidate.get("phone") or "").strip()
    email = (candidate.get("email") or "").strip()
    if redact_phone:
        if phone:
            digits = re.sub(r"\D", "", phone)
            if len(digits) >= 6:
                out = re.sub(re.escape(digits), "[REDACTED]", out)
        out = re.sub(r"(\+?\d[\d\s\-()]{8,}\d)", "[REDACTED]", out)
    if redact_email:
        if email:
            out = out.replace(email, "[REDACTED]")
        out = re.sub(r"[\w.+-]+@[\w-]+\.[\w.-]+", "[REDACTED]", out)
    return out


def fmt_exp(months) -> str:
    if not months:
        return ""
    y, m = divmod(int(months), 12)
    if y and m:
        return f"{y}y {m}m"
    if y:
        return f"{y}y"
    return f"{m}m"


_FS_UNSAFE_RE = re.compile(r'[\\/:*?"<>|]')


def build_resume_filename(display_name: str, position: Optional[str], total_exp_mo, ext: str) -> str:
    """"Candidate Name_Position_TotalExp.ext" — e.g. "Usha N_SAP FICO
    Consultant_12Yrs.pdf" — the one shared naming convention every resume-
    generation surface in the app uses (the standalone Resume Generator,
    the KAE/client-submission attachments, Standard Resume). Only
    genuinely filesystem-illegal characters are stripped; spaces WITHIN
    each field are kept (only underscores separate the 3 fields
    themselves), matching the reference format exactly rather than the
    older convention elsewhere in this codebase that replaced every space
    too. display_name already reflects masking (mask_name()) when that
    name_format is selected, so the filename stays consistent with the
    document's own header. A field that has no real value (no
    designation on file, no experience recorded) is simply omitted, never
    rendered as a blank segment."""
    def _clean(s):
        return _FS_UNSAFE_RE.sub("", (s or "").strip())
    name = _clean(display_name) or "Candidate"
    pos = _clean(position)
    exp = f"{round(total_exp_mo / 12)}Yrs" if total_exp_mo else ""
    parts = [p for p in (name, pos, exp) if p]
    return "_".join(parts) + f".{ext}"


DEFAULT_CONFIG = {
    "name_format": "full",          # full | masked
    "show_mobile": True,
    "show_email": True,
    "show_location": True,
    "company_mode": "original",     # original | replace | hide
    "company_replacement": None,
    "project_mode": "include",      # include | hide | focus
    "client_name_mode": "hide",     # show | hide | replace
    "client_name_replacement": None,
    "visual_theme": "classic",      # classic | modern_sidebar | minimal_ats
    "logo_position": "top_right",   # none | top_left | top_right
}

# Real improvement (2026-08-18, round 2): the logo was first added to the
# document FOOTER (replacing an earlier fixed "Generated via AVIIN ATS"
# text line) -- moved to a real page-HEADER placement per direct user
# feedback, with a genuine left/right choice rather than one fixed spot.
# "none" renders no logo at all. Same AVIIN Tech logo asset + sizing-by-
# real-aspect-ratio convention already established for call letters/offer
# letters (PDF via reportlab's Image flowable with hAlign; DOCX via
# python-docx's Run.add_picture with the containing paragraph's
# alignment). The "Submitted for: <client>" line (job-specific generation
# context, unrelated to logo placement) stays in the footer, unaffected.
LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "aviintech-logo.png")

# Real, distinct visual layouts a recruiter can pick between — separate
# dimension from the content-composition fields above (a "Project-Focused"
# template can render in any of these 3 looks). Not "clone an uploaded
# sample PDF" (that would need vision/AI analysis this zero-token codebase
# doesn't have) — a curated set of real layouts covering the styles
# recruiters actually asked for: a colored two-column layout (like a
# client-facing "showcase" resume) and a plain, color-free, single-column
# layout (the safest shape for automated ATS parsers on the client side).
VISUAL_THEMES = [
    {"id": "classic", "label": "Classic Professional",
     "description": "Centered header, single column, a blue accent rule. The current default look."},
    {"id": "modern_sidebar", "label": "Modern Sidebar",
     "description": "A shaded left sidebar for contact details and key skills, main column for the summary. A distinctive, client-facing look."},
    {"id": "minimal_ats", "label": "Minimal / ATS-Safe",
     "description": "Plain black text, no color, no tables, left-aligned. Optimized for automated resume-parsing systems."},
    # Real improvement (2026-08-18, round 4): 5 more real, distinct layouts
    # covering the styles seen across popular ATS-friendly and Canva-style
    # resume builders -- each built on the same safe pattern established
    # for the other themes today: any colored/shaded element is a small,
    # FIXED-size header block (never grows with resume length), while the
    # actual body always flows as plain paragraphs with no length cap, so
    # every one of these genuinely supports a full, unbounded-length
    # resume across as many pages as it needs.
    {"id": "executive_header", "label": "Executive Header",
     "description": "A bold solid-color header band with the name and title in white, plain single-column body below. A common senior/executive resume look."},
    {"id": "two_tone_header", "label": "Two-Tone Professional",
     "description": "Name on the left, a shaded contact panel on the right, single-column body below. A widely-used modern professional layout."},
    {"id": "timeline", "label": "Timeline Professional",
     "description": "Each role is marked with a colored bullet and a divider line, giving a career-timeline feel while staying fully ATS-parseable."},
    {"id": "compact_grid", "label": "Compact Professional",
     "description": "A skills grid strip up top and tighter body spacing -- a denser, more information-forward layout."},
    {"id": "elegant_serif", "label": "Elegant Serif",
     "description": "Refined serif typography with a double-rule header -- a premium, editorial look."},
]
_VALID_THEMES = {t["id"] for t in VISUAL_THEMES}

LOGO_POSITION_OPTIONS = [
    {"id": "top_left", "label": "Top Left", "description": "AVIIN Tech logo in the header, aligned left."},
    {"id": "top_right", "label": "Top Right", "description": "AVIIN Tech logo in the header, aligned right."},
    {"id": "none", "label": "No Logo", "description": "No logo anywhere in the document."},
]
_VALID_LOGO_POSITIONS = {o["id"] for o in LOGO_POSITION_OPTIONS}


_BULLET_ONLY_LINE_RE = re.compile(r'^[ \t]*[•❖◆●○■♦\-\*][ \t]*$', re.M)


def _strip_bullet_only_lines(text: str) -> str:
    """A source PDF's own text extraction sometimes splits 'bullet' and its
    sentence onto separate lines (a real PDF-layout artifact, not something
    this codebase introduced) -- e.g. a lone '•' line with the actual
    sentence appearing as its own paragraph right after. Rendered verbatim,
    that's a visibly empty bullet point on the generated resume. Real bug
    report (2026-08-18): a candidate's generated PDF showed two bare '•'
    lines with nothing after them. Drops any line that is ONLY a bullet
    character -- never touches a line that has real bullet+text together,
    which stays exactly as-is."""
    if not text:
        return text
    return _BULLET_ONLY_LINE_RE.sub('', text)


_MULTI_SPACE_RE = re.compile(r'[ \t]{2,}')
BULLET_LINE_RE = re.compile(r'^[ \t]*([•▪●○■♦])[ \t]*(.*)$')


def _normalize_whitespace(text: str) -> str:
    """Real improvement (2026-08-18): PDF text extraction (justified-text
    layouts especially) commonly leaves runs of 2+ spaces between words --
    confirmed directly in a real resume's stored text ("I  worked  in
    Technip  Energies", "•  Working with Burns..."). Rendered verbatim,
    that's visibly uneven spacing throughout the generated document, most
    noticeable right after a bullet character. Collapses any run of 2+
    spaces/tabs to exactly one -- never touches newlines, which are real
    line/paragraph boundaries, not extraction noise."""
    if not text:
        return text
    return '\n'.join(_MULTI_SPACE_RE.sub(' ', line) for line in text.split('\n'))


_SENTENCE_END_RE = re.compile(r'[.:;!?]$')


def _merge_wrapped_lines(text: str) -> list[str]:
    """Real fix (2026-08-18): confirmed directly against this candidate's
    own stored resume_text that the source PDF's text extraction hard-wraps
    long lines at a fixed character width, independent of sentence
    boundaries -- e.g. "...Set up Network profiles and\\nsettlement
    profiles. Configured..." is really ONE continuous sentence, split into
    two raw newline-separated lines with no punctuation at the break.
    Splitting purely on '\\n' (every renderer's loop, until now) treated
    each wrapped fragment as its own paragraph -- harmless when rendered as
    plain body text, but actively wrong once _classify_lines() (added the
    same day) started auto-bulleting achievement-shaped lines: each
    fragment got its OWN bullet, turning one real achievement into several
    nonsensical ones. Reflows wrapped fragments back into logical lines
    before any classification runs: a line only starts a new logical line
    when it is itself a literal-bulleted or subheading line, the line
    before it is a subheading (headings never absorb the next line), or
    the line before it already ends in real sentence-terminal punctuation,
    or the line before it is itself too short to plausibly be a hard-wrap
    artifact -- otherwise it's a wrapped continuation and gets appended
    onto the previous logical line.

    Real bug caught and fixed while verifying this against the actual
    candidate above: without a minimum-length guard, a genuinely short,
    separate line with no trailing punctuation -- a role title
    ("SAP S4 HANA Finance Consultant") or a short standalone bullet
    ("Worked in Migrations") -- got wrongly absorbed into whatever
    followed it, since it looked identical to a real wrapped fragment by
    the punctuation rule alone. Every genuine wrapped fragment observed
    in this document's own real hard-wrap width ran 90-115 characters;
    60 is a conservative floor under that, so short lines stay standalone
    while long unterminated lines still merge correctly."""
    _WRAP_MIN_LEN = 60
    raw = [l.strip() for l in text.split('\n') if l.strip()]
    merged: list[str] = []
    for line in raw:
        starts_new = (
            not merged
            or bool(BULLET_LINE_RE.match(line))
            or _is_subheading(line)
            or _is_subheading(merged[-1])
            or bool(_SENTENCE_END_RE.search(merged[-1]))
            or len(merged[-1]) < _WRAP_MIN_LEN
        )
        if starts_new:
            merged.append(line)
        else:
            merged[-1] = f"{merged[-1]} {line}"
    return merged


def _classify_lines(text: str) -> list[tuple[str, str]]:
    """Real improvement (2026-08-18): many real resumes visually bullet
    every achievement/responsibility line under a role via the SOURCE
    PDF's own list formatting -- a layout property, not a literal
    character in the text stream -- so plain-text extraction silently
    drops that signal for any line that didn't happen to have a real '•'
    character. Confirmed directly against a real resume where an entire
    role's worth of achievement sentences ("Established internal order
    settlements...", "Implemented EBS reconciliation...") rendered with
    no bullet at all, reading as an undifferentiated wall of sentences,
    while OTHER roles nearby (which happened to have literal bullets in
    the extracted text) looked correct. Standard reverse-chronological
    resume convention bullets every real achievement sentence under a
    role while leaving the role's own title line and section headers
    alone -- this function reproduces that, conservatively:

    Returns [(display_text, kind), ...] where kind is 'subhead', 'bullet',
    or 'body'. A plain (non-already-bulleted, non-subheading) line only
    becomes an auto-bullet when ALL of: (a) we're inside a real
    PROFESSIONAL EXPERIENCE-style block -- seen a "professional
    experience"/"work experience"/"employment history" heading, OR a
    real "Employer:" line (a resume-specific signal just as strong even
    with no generic heading present); (b) it is NOT the line immediately
    after an "Employer:"/"Client:" line (that slot is almost always the
    role's own title, e.g. "SAP FICO Lead Constant" -- never an
    achievement); (c) it reads like a real sentence (ends with '.' or is
    a genuinely long line). Deliberately conservative in the same spirit
    as _is_subheading() above -- never applied outside a real experience
    block, so the PROFESSIONAL SUMMARY's own intro paragraph and other
    prose stay untouched. Known, accepted limitation: this resume's own
    source structure inconsistently places a role's title line after
    "Employer:" for some roles and after "Client:" for others -- when a
    title line is genuinely absent in a spot this heuristic expects one,
    the very next real achievement line is conservatively left
    unbulleted rather than risk a title line wrongly gaining a bullet."""
    out: list[tuple[str, str]] = []
    in_experience = False
    prev_was_role_header = False
    for p in _merge_wrapped_lines(text):
        bm = BULLET_LINE_RE.match(p)
        if bm:
            out.append((bm.group(2).strip(), 'bullet'))
            prev_was_role_header = False
            continue
        if _is_subheading(p):
            normalized = p.rstrip(':').strip().lower()
            low = p.lower()
            is_role_header = low.startswith('employer:') or low.startswith('client:')
            if normalized in ('professional experience', 'work experience', 'employment history') or is_role_header:
                in_experience = True
            out.append((p, 'subhead'))
            prev_was_role_header = is_role_header
            continue
        if in_experience and not prev_was_role_header and (p.endswith('.') or len(p) > 40):
            out.append((p, 'bullet'))
        else:
            out.append((p, 'body'))
        prev_was_role_header = False
    return out


def _resolve_body_text(candidate: dict, config: dict) -> tuple[str, str]:
    """Returns (section_heading, body_text) for the main narrative block,
    already redacted/masked per config. project_mode='focus' isolates just
    the Projects section (falls back to the full summary with a note if the
    resume has no distinct Projects heading); 'hide' omits the narrative
    entirely; 'include' is the normal full extracted text."""
    from services.improved_parser import extract_projects_section, extract_summary_section

    resume_text = candidate.get("resume_text") or ""
    if config["project_mode"] == "hide":
        return "", ""

    if config["project_mode"] == "focus":
        projects = extract_projects_section(resume_text)
        if not projects:
            return "", ""
        heading = "PROJECTS"
        raw = projects
    else:
        # REAL BUG FIX (2026-08-18): this used to be `raw = resume_text`
        # -- the *entire* raw document, including its own embedded name/
        # title/"PROFESSIONAL SUMMARY" heading -- rendered under a
        # second, template-added "PROFESSIONAL SUMMARY" heading. Every
        # generated resume showed the candidate's name and title twice.
        heading = "PROFESSIONAL SUMMARY"
        raw = extract_summary_section(resume_text, candidate.get("full_name") or "") or ""

    text = redact_contact(raw, candidate, not config["show_mobile"], not config["show_email"])
    full_name = candidate.get("full_name") or ""
    if config["name_format"] == "masked" and full_name:
        text = re.sub(re.escape(full_name), mask_name(full_name), text, flags=re.I)
    employer = candidate.get("current_employer") or ""
    if employer and config["company_mode"] == "replace" and config["company_replacement"]:
        text = re.sub(re.escape(employer), config["company_replacement"], text, flags=re.I)
    elif employer and config["company_mode"] == "hide":
        text = re.sub(re.escape(employer), "[Company withheld]", text, flags=re.I)
    text = _strip_bullet_only_lines(text)
    text = _normalize_whitespace(text)
    return heading, text


def _company_line(candidate: dict, config: dict) -> Optional[str]:
    employer = candidate.get("current_employer") or ""
    if config["company_mode"] == "hide":
        return None
    if config["company_mode"] == "replace":
        return config["company_replacement"] or employer or None
    return employer or None


def _client_line(client_name: Optional[str], config: dict) -> Optional[str]:
    if not client_name or config["client_name_mode"] == "hide":
        return None
    if config["client_name_mode"] == "replace":
        return config["client_name_replacement"] or client_name
    return client_name


def _pdf_header_logo_flowables(cfg: dict) -> list:
    """Real improvement (2026-08-18, round 2): the logo moved from the
    footer to the page header per direct user feedback, with a genuine
    left/right placement choice -- reportlab's Image flowable already
    supports hAlign natively, so a real top-left/top-right header is just
    the image as the very first flowable with hAlign set accordingly (no
    absolute positioning/canvas tricks needed). Returns [] for
    logo_position="none" or a missing asset file."""
    pos = cfg.get("logo_position", "top_right")
    if pos not in ("top_left", "top_right") or not os.path.exists(LOGO_PATH):
        return []
    from reportlab.lib.units import cm
    from reportlab.platypus import Image, Spacer
    logo_w = 2.6 * cm
    logo_h = logo_w * (342 / 730)
    img = Image(LOGO_PATH, width=logo_w, height=logo_h)
    img.hAlign = "LEFT" if pos == "top_left" else "RIGHT"
    return [img, Spacer(1, 0.25 * cm)]


def _docx_header_logo(doc, cfg: dict) -> None:
    """DOCX analogue of _pdf_header_logo_flowables above -- python-docx
    positions an inline picture via the containing paragraph's own
    alignment (LEFT/RIGHT), inserted as the very first paragraph in the
    document."""
    pos = cfg.get("logo_position", "top_right")
    if pos not in ("top_left", "top_right") or not os.path.exists(LOGO_PATH):
        return
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if pos == "top_left" else WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run().add_picture(LOGO_PATH, width=Cm(2.6))


def _pdf_footer_flowables(cfg: dict, client_line: Optional[str], small_style) -> list:
    """The footer now only ever carries the "Submitted for: <client>"
    line (job-specific generation context) -- the logo lives in the
    header, see _pdf_header_logo_flowables above."""
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, Spacer
    if not client_line:
        return []
    return [Spacer(1, 0.5 * cm), Paragraph(_esc(f"Submitted for: {client_line}"), small_style)]


def _docx_footer(doc, cfg: dict, client_line: Optional[str], *, italic: bool = True,
                  centered: bool = True, size: float = 9, color=None) -> None:
    """DOCX analogue of _pdf_footer_flowables above -- text-only now, the
    logo lives in the header (_docx_header_logo)."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if client_line:
        fp = doc.add_paragraph()
        if centered:
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(f"Submitted for: {client_line}")
        fr.italic = italic
        fr.font.size = Pt(size)
        if color is not None:
            fr.font.color.rgb = color


_PDF_RENDERERS = {}  # populated after all _render_pdf_* functions are defined, see bottom of file
_DOCX_RENDERERS = {}  # populated after all _render_docx_* functions are defined, see bottom of file


def render_resume_pdf(candidate: dict, config: dict, client_name: str = None) -> bytes:
    """Dispatches to one of VISUAL_THEMES's real renderers. All of them use
    the exact same resolved content (_resolve_body_text/_company_line/etc)
    — only layout, color, and typography differ."""
    cfg = {**DEFAULT_CONFIG, **config}
    theme = cfg.get("visual_theme") or "classic"
    if theme not in _VALID_THEMES:
        theme = "classic"
    return _PDF_RENDERERS.get(theme, _render_pdf_classic)(candidate, cfg, client_name)


def _render_pdf_classic(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                             topMargin=2 * cm, bottomMargin=2 * cm)
    PRIMARY = colors.HexColor("#1e40af")
    DARK = colors.HexColor("#0f172a")
    GRAY = colors.HexColor("#64748b")
    # REAL BUG FIX (2026-08-18): reportlab's ParagraphStyle defaults
    # `leading` (line height) to a FIXED 12pt regardless of fontSize --
    # confirmed directly (`ParagraphStyle(fontSize=18).leading == 12`).
    # Every heading style below was larger than 12pt with no explicit
    # leading, so the text rendered taller than its allocated line box and
    # visually overlapped the paragraph right after it -- confirmed on a
    # real generated PDF (name and title overlapping at the very top of
    # the page). This has been true since the renderer was first built;
    # only surfaced now because this is the first time a generated PDF was
    # actually rendered to an image and inspected, rather than checked via
    # text extraction (which reveals content, not visual layout).
    h1 = ParagraphStyle("H1", fontSize=18, leading=22, textColor=DARK, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)
    sub = ParagraphStyle("Sub", fontSize=11, leading=14, textColor=PRIMARY, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=10)
    h2 = ParagraphStyle("H2", fontSize=11, leading=14, textColor=PRIMARY, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("Body", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica")
    # Real improvement (2026-08-18): bold in-body sub-headings (section
    # titles, "Employer:"/"Client:" lines) -- same size/color as body text
    # so it reads as emphasis, not a competing heading level against h2.
    subhead = ParagraphStyle("SubHead", fontSize=10, leading=15, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=2)
    small = ParagraphStyle("Small", fontSize=9, leading=11, textColor=GRAY, fontName="Helvetica")
    # Real improvement (2026-08-18): reportlab's `bulletText` param gives a
    # real hanging indent -- a wrapped bullet line's second line aligns
    # under the FIRST WORD after the bullet, not back at the left margin,
    # matching how a real bulleted list looks (confirmed against the
    # source document's own rendering). Rendering "• text" as one plain
    # string with no indent (the previous approach) left wrapped lines
    # flush left, visually indistinguishable from a new paragraph.
    bullet = ParagraphStyle("Bullet", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica",
                             leftIndent=14, bulletIndent=0, spaceBefore=1, spaceAfter=1)

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    story = _pdf_header_logo_flowables(cfg) + [
        Paragraph(_esc(display_name), h1),
        Paragraph(_esc(candidate.get("current_designation") or ""), sub),
        HRFlowable(width="100%", thickness=1.2, color=PRIMARY, spaceAfter=6),
    ]

    meta_bits = []
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"<b>Location:</b> {_esc(candidate['location'])}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"<b>Total Experience:</b> {_esc(fmt_exp(candidate['total_exp_mo']))}")
    if meta_bits:
        story.append(Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(meta_bits), body))

    contact_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        contact_bits.append(f"<b>Mobile:</b> {_esc(candidate['phone'])}")
    if cfg["show_email"] and candidate.get("email"):
        contact_bits.append(f"<b>Email:</b> {_esc(candidate['email'])}")
    if contact_bits:
        story.append(Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(contact_bits), body))

    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"<b>Current Company:</b> {_esc(company_line)}", body))

    skills = candidate.get("skills") or []
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(", ".join(skills)), body))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        # REAL BUG FIX (2026-08-18): this used to hard-cap the rendered
        # body at 2600 chars regardless of how much real content
        # _resolve_body_text() actually returned -- for a dense, multi-
        # role resume, the entire Professional Experience/Education/
        # Certifications section (everything past the opening summary
        # paragraph) was silently cut off with an ellipsis, even after
        # today's earlier fix restored the full text into the pipeline.
        # SimpleDocTemplate/python-docx both paginate naturally across as
        # many pages as the real content needs -- no reason to
        # artificially truncate before handing it to them.
        # Real improvement (2026-08-18): _classify_lines() auto-bullets real
        # achievement sentences under a role that never had a literal bullet
        # character in the extracted text -- a common case, since many
        # source resumes convey their bullets as a layout property of the
        # original document rather than a character in the text stream --
        # in addition to preserving lines that already had a real one.
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(_esc(line_text), subhead))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def _render_pdf_sidebar(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """A real 2-column layout on page 1: a shaded left sidebar (contact +
    key skills), a white main column (name/title header + the resolved
    summary/projects body). Real fix (2026-08-18, round 3): this used to be
    ONE reportlab Table row -- which cannot split across pages -- so any
    resume too long for a single page silently got hard-truncated at a
    fixed character count with a trailing "…", regardless of how much real
    content the candidate actually had. Direct user report: this reads as
    broken, not "intentionally condensed." Rebuilt on BaseDocTemplate with
    two real PageTemplates instead of a single Table: page 1 keeps the
    colored sidebar (drawn via onPage, not a Table cell background) next to
    a main-column Frame; a FrameBreak() moves from the sidebar frame into
    the main frame, and a NextPageTemplate('Continuation') queued right
    after it means any content that overflows page 1's main frame flows
    onto plain, full-width, sidebar-free continuation pages -- the same
    "cover page then flowing content" idiom reportlab's own docs use for
    exactly this case. No content is ever cut short and no page count is
    capped; a short resume still renders as a clean single page, a long one
    genuinely spans as many pages as it needs, matching how the classic and
    minimal themes already behave."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, FrameBreak,
                                     NextPageTemplate, Paragraph, Spacer)

    buf = io.BytesIO()
    page_w, page_h = A4
    sidebar_w = page_w * 0.32
    main_w = page_w - sidebar_w
    bottom_margin = 1.2 * cm

    SIDEBAR_BG = colors.HexColor("#1e3a5f")
    SIDEBAR_TEXT = colors.HexColor("#e8eef5")
    SIDEBAR_ACCENT = colors.HexColor("#7fb3e0")
    DARK = colors.HexColor("#0f172a")
    PRIMARY = colors.HexColor("#1e40af")
    GRAY = colors.HexColor("#64748b")

    sb_label = ParagraphStyle("SbLabel", fontSize=9, leading=11, textColor=SIDEBAR_ACCENT, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
    sb_body = ParagraphStyle("SbBody", fontSize=9, textColor=SIDEBAR_TEXT, leading=13, fontName="Helvetica")
    m_name = ParagraphStyle("MName", fontSize=20, leading=24, textColor=DARK, fontName="Helvetica-Bold", spaceAfter=4)
    m_title = ParagraphStyle("MTitle", fontSize=12, leading=15, textColor=PRIMARY, fontName="Helvetica-Bold", spaceAfter=14)
    m_h2 = ParagraphStyle("MH2", fontSize=11, leading=14, textColor=PRIMARY, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=6)
    m_body = ParagraphStyle("MBody", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica")
    m_subhead = ParagraphStyle("MSubHead", fontSize=10, leading=15, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=2)
    m_bullet = ParagraphStyle("MBullet", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica",
                               leftIndent=12, bulletIndent=0, spaceBefore=1, spaceAfter=1)
    small = ParagraphStyle("Small", fontSize=8, leading=10, textColor=GRAY, fontName="Helvetica")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")

    # ── Sidebar frame content (page 1 only, fixed-length -- contact/skills
    # never realistically overflow a page on their own) ──
    sidebar = [Spacer(1, 1.2 * cm)]
    sidebar.append(Paragraph("CONTACT", sb_label))
    if cfg["show_mobile"] and candidate.get("phone"):
        sidebar.append(Paragraph(_esc(candidate["phone"]), sb_body))
    if cfg["show_email"] and candidate.get("email"):
        sidebar.append(Paragraph(_esc(candidate["email"]), sb_body))
    if cfg["show_location"] and candidate.get("location"):
        sidebar.append(Paragraph(_esc(candidate["location"]), sb_body))
    if candidate.get("total_exp_mo"):
        sidebar.append(Paragraph("EXPERIENCE", sb_label))
        sidebar.append(Paragraph(_esc(fmt_exp(candidate["total_exp_mo"])), sb_body))
    company_line = _company_line(candidate, cfg)
    if company_line:
        sidebar.append(Paragraph("CURRENT COMPANY", sb_label))
        sidebar.append(Paragraph(_esc(company_line), sb_body))
    skills = candidate.get("skills") or []
    if skills:
        sidebar.append(Paragraph("KEY SKILLS", sb_label))
        for s in skills[:18]:
            sidebar.append(Paragraph(f"• {_esc(s)}", sb_body))

    # ── Main column content -- real header logo (inline, not a separate
    # Table row, since it now lives inside the flowing main frame) + name/
    # title + the FULL resolved body, no length cap. ──
    header_logo = _pdf_header_logo_flowables(cfg)
    if header_logo:
        main = [Spacer(1, 0.8 * cm)] + header_logo + [Spacer(1, 0.2 * cm)]
    else:
        main = [Spacer(1, 1.2 * cm)]
    main.append(Paragraph(_esc(display_name), m_name))
    if candidate.get("current_designation"):
        main.append(Paragraph(_esc(candidate["current_designation"]), m_title))
    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        main.append(Paragraph(heading, m_h2))
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                main.append(Paragraph(_esc(line_text), m_bullet, bulletText='•'))
            elif kind == 'subhead':
                main.append(Paragraph(_esc(line_text), m_subhead))
            else:
                main.append(Paragraph(_esc(line_text), m_body))
    client_line = _client_line(client_name, cfg)
    main.extend(_pdf_footer_flowables(cfg, client_line, small))

    def _draw_sidebar_bg(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(SIDEBAR_BG)
        canvas.rect(0, 0, sidebar_w, page_h, fill=1, stroke=0)
        canvas.restoreState()

    sidebar_frame = Frame(0, 0, sidebar_w, page_h, leftPadding=0.9 * cm, rightPadding=0.7 * cm,
                           topPadding=0, bottomPadding=0, id="sidebar", showBoundary=0)
    main_frame = Frame(sidebar_w, bottom_margin, main_w, page_h - bottom_margin,
                        leftPadding=1.2 * cm, rightPadding=1.2 * cm, topPadding=0, bottomPadding=0,
                        id="main", showBoundary=0)
    cont_frame = Frame(2.2 * cm, 2 * cm, page_w - 4.4 * cm, page_h - 4 * cm, id="cont", showBoundary=0)

    doc = BaseDocTemplate(buf, pagesize=A4)
    doc.addPageTemplates([
        PageTemplate(id="First", frames=[sidebar_frame, main_frame], onPage=_draw_sidebar_bg),
        PageTemplate(id="Continuation", frames=[cont_frame]),
    ])

    # Paragraph flowables wrap dynamically against whatever frame they end
    # up rendered into -- the same m_body/m_bullet/m_subhead styles used
    # above read correctly whether a given paragraph lands in page 1's
    # narrower main frame or a wider continuation frame; only the wrap
    # width changes, not the font metrics.
    story = sidebar + [FrameBreak(), NextPageTemplate("Continuation")] + main
    doc.build(story)
    return buf.getvalue()


def _render_pdf_minimal(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """No color, no tables, no centering — the shape most likely to survive
    an automated ATS text-extraction pass on the client's own side."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.4 * cm, rightMargin=2.4 * cm,
                             topMargin=2.2 * cm, bottomMargin=2.2 * cm)
    BLACK = colors.HexColor("#000000")
    GRAY = colors.HexColor("#444444")
    # REAL BUG FIX (2026-08-18): see the identical fix + explanation in
    # _render_pdf_classic above.
    h1 = ParagraphStyle("H1", fontSize=16, leading=19, textColor=BLACK, fontName="Helvetica-Bold", spaceAfter=4)
    sub = ParagraphStyle("Sub", fontSize=11, leading=14, textColor=BLACK, fontName="Helvetica", spaceAfter=8)
    h2 = ParagraphStyle("H2", fontSize=10.5, leading=13, textColor=BLACK, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=5)
    body = ParagraphStyle("Body", fontSize=10, textColor=BLACK, leading=14, fontName="Helvetica")
    # Real improvement (2026-08-18): bold in-body sub-headings, matching
    # the identical addition in _render_pdf_classic above.
    subhead = ParagraphStyle("SubHead", fontSize=10, leading=14, textColor=BLACK, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=2)
    # Real improvement (2026-08-18): hanging-indent bullets, matching the
    # identical addition in _render_pdf_classic above.
    bullet = ParagraphStyle("Bullet", fontSize=10, textColor=BLACK, leading=14, fontName="Helvetica",
                             leftIndent=14, bulletIndent=0, spaceBefore=1, spaceAfter=1)
    small = ParagraphStyle("Small", fontSize=8.5, leading=10.5, textColor=GRAY, fontName="Helvetica")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    story = _pdf_header_logo_flowables(cfg) + [Paragraph(_esc(display_name), h1)]
    if candidate.get("current_designation"):
        story.append(Paragraph(_esc(candidate["current_designation"]), sub))
    story.append(HRFlowable(width="100%", thickness=0.6, color=GRAY, spaceAfter=8))

    line_bits = []
    if cfg["show_location"] and candidate.get("location"):
        line_bits.append(candidate["location"])
    if candidate.get("total_exp_mo"):
        line_bits.append(f"{fmt_exp(candidate['total_exp_mo'])} experience")
    if cfg["show_mobile"] and candidate.get("phone"):
        line_bits.append(candidate["phone"])
    if cfg["show_email"] and candidate.get("email"):
        line_bits.append(candidate["email"])
    if line_bits:
        story.append(Paragraph(_esc(" | ".join(line_bits)), body))

    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"Current Company: {_esc(company_line)}", body))

    skills = candidate.get("skills") or []
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(", ".join(skills)), body))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        # REAL BUG FIX (2026-08-18): this used to hard-cap the rendered
        # body at 2600 chars regardless of how much real content
        # _resolve_body_text() actually returned -- for a dense, multi-
        # role resume, the entire Professional Experience/Education/
        # Certifications section (everything past the opening summary
        # paragraph) was silently cut off with an ellipsis, even after
        # today's earlier fix restored the full text into the pipeline.
        # SimpleDocTemplate/python-docx both paginate naturally across as
        # many pages as the real content needs -- no reason to
        # artificially truncate before handing it to them.
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(_esc(line_text), subhead))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def render_resume_docx(candidate: dict, config: dict, client_name: str = None) -> bytes:
    """Dispatches to one of VISUAL_THEMES's real DOCX renderers, mirroring
    the PDF dispatcher above."""
    cfg = {**DEFAULT_CONFIG, **config}
    theme = cfg.get("visual_theme") or "classic"
    if theme not in _VALID_THEMES:
        theme = "classic"
    return _DOCX_RENDERERS.get(theme, _render_docx_classic)(candidate, cfg, client_name)


def _render_docx_classic(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    PRIMARY = RGBColor(0x1e, 0x40, 0xaf)
    DARK = RGBColor(0x0f, 0x17, 0x2a)

    doc = Document()
    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")

    _docx_header_logo(doc, cfg)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK

    if candidate.get("current_designation"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(candidate["current_designation"])
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = PRIMARY

    meta_bits = []
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"Location: {candidate['location']}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"Total Experience: {fmt_exp(candidate['total_exp_mo'])}")
    if meta_bits:
        doc.add_paragraph(" | ".join(meta_bits))

    contact_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        contact_bits.append(f"Mobile: {candidate['phone']}")
    if cfg["show_email"] and candidate.get("email"):
        contact_bits.append(f"Email: {candidate['email']}")
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    company_line = _company_line(candidate, cfg)
    if company_line:
        doc.add_paragraph(f"Current Company: {company_line}")

    skills = candidate.get("skills") or []
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.color.rgb = PRIMARY
        doc.add_paragraph(", ".join(skills))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = PRIMARY
        # REAL BUG FIX (2026-08-18): this used to hard-cap the rendered
        # body at 2600 chars regardless of how much real content
        # _resolve_body_text() actually returned -- for a dense, multi-
        # role resume, the entire Professional Experience/Education/
        # Certifications section (everything past the opening summary
        # paragraph) was silently cut off with an ellipsis, even after
        # today's earlier fix restored the full text into the pipeline.
        # SimpleDocTemplate/python-docx both paginate naturally across as
        # many pages as the real content needs -- no reason to
        # artificially truncate before handing it to them.
        # Real improvement (2026-08-18): a real bulleted paragraph (Word's
        # built-in "List Bullet" style, available by default -- no custom
        # numbering XML needed) instead of a literal "• " prefix on a plain
        # paragraph, which had no hanging indent -- a wrapped bullet line
        # fell back to the left margin instead of aligning under the text.
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                doc.add_paragraph(line_text, style='List Bullet')
                continue
            pp = doc.add_paragraph()
            r = pp.add_run(line_text)
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = DARK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _shade_cell(cell, hex_color: str) -> None:
    """python-docx has no first-class cell-shading API — the standard
    workaround is a raw <w:shd> element on the cell's tcPr, same technique
    used throughout the python-docx ecosystem for table-cell backgrounds."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _render_docx_sidebar(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """A single-row, 2-column table: shaded left cell (contact + skills),
    white right cell (name/title + resolved summary) — the DOCX analogue of
    the PDF sidebar theme's reportlab Table approach."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    SIDEBAR_TEXT = RGBColor(0xe8, 0xee, 0xf5)
    SIDEBAR_ACCENT = RGBColor(0x7f, 0xb3, 0xe0)
    PRIMARY = RGBColor(0x1e, 0x40, 0xaf)
    DARK = RGBColor(0x0f, 0x17, 0x2a)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    _docx_header_logo(doc, cfg)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(6)
    right.width = Cm(13)
    _shade_cell(left, "1E3A5F")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")

    def sb_heading(txt):
        p = left.add_paragraph()
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = SIDEBAR_ACCENT

    def sb_line(txt):
        p = left.add_paragraph()
        r = p.add_run(txt)
        r.font.size = Pt(9)
        r.font.color.rgb = SIDEBAR_TEXT

    left.paragraphs[0].text = ""
    sb_heading("CONTACT")
    if cfg["show_mobile"] and candidate.get("phone"):
        sb_line(candidate["phone"])
    if cfg["show_email"] and candidate.get("email"):
        sb_line(candidate["email"])
    if cfg["show_location"] and candidate.get("location"):
        sb_line(candidate["location"])
    if candidate.get("total_exp_mo"):
        sb_heading("EXPERIENCE")
        sb_line(fmt_exp(candidate["total_exp_mo"]))
    company_line = _company_line(candidate, cfg)
    if company_line:
        sb_heading("CURRENT COMPANY")
        sb_line(company_line)
    skills = candidate.get("skills") or []
    if skills:
        sb_heading("KEY SKILLS")
        for s in skills[:18]:
            sb_line(f"• {s}")

    right.paragraphs[0].text = ""
    p = right.paragraphs[0]
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK

    if candidate.get("current_designation"):
        p = right.add_paragraph()
        r = p.add_run(candidate["current_designation"])
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = PRIMARY

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = right.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = PRIMARY
        # REAL FIX (2026-08-18, round 3): no longer capped. Unlike a
        # reportlab PDF Table (which cannot split a row across pages, the
        # reason the PDF sidebar renderer was rebuilt on real multi-page
        # Frames this same round), a python-docx table's row is allowed to
        # break across pages in Word by default (no cantSplit set) -- so
        # this single-row 2-column table already handles long content
        # correctly without any special-casing, matching the PDF version's
        # new real-pagination behavior instead of silently truncating.
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                right.add_paragraph(line_text, style='List Bullet')
                continue
            pp = right.add_paragraph()
            r = pp.add_run(line_text)
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = DARK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_minimal(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """No color, left-aligned, no italics — the DOCX analogue of the PDF
    minimal/ATS-safe theme."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    BLACK = RGBColor(0x00, 0x00, 0x00)
    GRAY = RGBColor(0x44, 0x44, 0x44)

    doc = Document()
    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")

    _docx_header_logo(doc, cfg)

    p = doc.add_paragraph()
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = BLACK

    if candidate.get("current_designation"):
        p = doc.add_paragraph()
        r = p.add_run(candidate["current_designation"])
        r.font.size = Pt(11)
        r.font.color.rgb = BLACK

    line_bits = []
    if cfg["show_location"] and candidate.get("location"):
        line_bits.append(candidate["location"])
    if candidate.get("total_exp_mo"):
        line_bits.append(f"{fmt_exp(candidate['total_exp_mo'])} experience")
    if cfg["show_mobile"] and candidate.get("phone"):
        line_bits.append(candidate["phone"])
    if cfg["show_email"] and candidate.get("email"):
        line_bits.append(candidate["email"])
    if line_bits:
        doc.add_paragraph(" | ".join(line_bits))

    company_line = _company_line(candidate, cfg)
    if company_line:
        doc.add_paragraph(f"Current Company: {company_line}")

    skills = candidate.get("skills") or []
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.color.rgb = BLACK
        doc.add_paragraph(", ".join(skills))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = BLACK
        # REAL BUG FIX (2026-08-18): this used to hard-cap the rendered
        # body at 2600 chars regardless of how much real content
        # _resolve_body_text() actually returned -- for a dense, multi-
        # role resume, the entire Professional Experience/Education/
        # Certifications section (everything past the opening summary
        # paragraph) was silently cut off with an ellipsis, even after
        # today's earlier fix restored the full text into the pipeline.
        # SimpleDocTemplate/python-docx both paginate naturally across as
        # many pages as the real content needs -- no reason to
        # artificially truncate before handing it to them.
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                doc.add_paragraph(line_text, style='List Bullet')
                continue
            pp = doc.add_paragraph()
            r = pp.add_run(line_text)
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = BLACK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line, italic=False, centered=False, size=8.5, color=GRAY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────── 5 additional themes (round 4) ───────────────────────
# Every one of these follows the same safety rule established across today's
# earlier fixes: any colored/shaded element is a small, FIXED-size flowable
# (a short header Table, a skills grid capped the same way the sidebar theme
# already caps it) that never grows with resume length, while the actual
# body always flows as plain uncapped paragraphs -- so all 5 genuinely
# support a full, unbounded-length resume across as many pages as it needs,
# the same guarantee the classic/minimal/sidebar themes already have.

def _render_pdf_executive_header(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """A bold solid-color header band (name + title in white) followed by a
    plain single-column ATS-parseable body -- a common senior/executive
    resume look."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                             topMargin=1.6 * cm, bottomMargin=2 * cm)
    BAND = colors.HexColor("#1e3a5f")
    PRIMARY = colors.HexColor("#1e40af")
    DARK = colors.HexColor("#0f172a")
    GRAY = colors.HexColor("#64748b")

    h1_white = ParagraphStyle("H1W", fontSize=21, leading=25, textColor=colors.white, fontName="Helvetica-Bold")
    sub_white = ParagraphStyle("SubW", fontSize=12, leading=15, textColor=colors.HexColor("#cbd5e1"), fontName="Helvetica", spaceBefore=2)
    h2 = ParagraphStyle("H2", fontSize=11, leading=14, textColor=PRIMARY, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("Body", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica")
    subhead = ParagraphStyle("SubHead", fontSize=10, leading=15, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=2)
    bullet = ParagraphStyle("Bullet", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica",
                             leftIndent=14, bulletIndent=0, spaceBefore=1, spaceAfter=1)
    small = ParagraphStyle("Small", fontSize=9, leading=11, textColor=GRAY, fontName="Helvetica")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    band_content = [Paragraph(_esc(display_name), h1_white)]
    if candidate.get("current_designation"):
        band_content.append(Paragraph(_esc(candidate["current_designation"]), sub_white))
    band = Table([[band_content]], colWidths=[doc.width])
    band.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), BAND),
        ("LEFTPADDING", (0, 0), (-1, -1), 0.8 * cm), ("RIGHTPADDING", (0, 0), (-1, -1), 0.8 * cm),
        ("TOPPADDING", (0, 0), (-1, -1), 0.7 * cm), ("BOTTOMPADDING", (0, 0), (-1, -1), 0.7 * cm),
    ]))

    story = _pdf_header_logo_flowables(cfg) + [band, Spacer(1, 0.5 * cm)]

    meta_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        meta_bits.append(f"<b>Mobile:</b> {_esc(candidate['phone'])}")
    if cfg["show_email"] and candidate.get("email"):
        meta_bits.append(f"<b>Email:</b> {_esc(candidate['email'])}")
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"<b>Location:</b> {_esc(candidate['location'])}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"<b>Experience:</b> {_esc(fmt_exp(candidate['total_exp_mo']))}")
    if meta_bits:
        story.append(Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(meta_bits), body))

    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"<b>Current Company:</b> {_esc(company_line)}", body))

    skills = candidate.get("skills") or []
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(", ".join(skills)), body))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(_esc(line_text), subhead))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def _render_pdf_two_tone_header(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """Name on the left, a shaded contact panel on the right, single-column
    body below -- a widely-used modern professional layout."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                             topMargin=2 * cm, bottomMargin=2 * cm)
    SHADE = colors.HexColor("#eff6ff")
    PRIMARY = colors.HexColor("#1e40af")
    DARK = colors.HexColor("#0f172a")
    GRAY = colors.HexColor("#64748b")

    h1 = ParagraphStyle("H1", fontSize=19, leading=23, textColor=DARK, fontName="Helvetica-Bold")
    sub = ParagraphStyle("Sub", fontSize=11.5, leading=14, textColor=PRIMARY, fontName="Helvetica-Bold", spaceBefore=2)
    contact = ParagraphStyle("Contact", fontSize=9, leading=13, textColor=DARK, fontName="Helvetica")
    h2 = ParagraphStyle("H2", fontSize=11, leading=14, textColor=PRIMARY, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("Body", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica")
    subhead = ParagraphStyle("SubHead", fontSize=10, leading=15, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=2)
    bullet = ParagraphStyle("Bullet", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica",
                             leftIndent=14, bulletIndent=0, spaceBefore=1, spaceAfter=1)
    small = ParagraphStyle("Small", fontSize=9, leading=11, textColor=GRAY, fontName="Helvetica")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    left = [Paragraph(_esc(display_name), h1)]
    if candidate.get("current_designation"):
        left.append(Paragraph(_esc(candidate["current_designation"]), sub))

    right = []
    if cfg["show_mobile"] and candidate.get("phone"):
        right.append(Paragraph(_esc(candidate["phone"]), contact))
    if cfg["show_email"] and candidate.get("email"):
        right.append(Paragraph(_esc(candidate["email"]), contact))
    if cfg["show_location"] and candidate.get("location"):
        right.append(Paragraph(_esc(candidate["location"]), contact))
    if candidate.get("total_exp_mo"):
        right.append(Paragraph(_esc(fmt_exp(candidate["total_exp_mo"])) + " experience", contact))
    if not right:
        right = [Paragraph("", contact)]

    header_table = Table([[left, right]], colWidths=[doc.width * 0.62, doc.width * 0.38])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (1, 0), (1, 0), SHADE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (0, 0), 0), ("RIGHTPADDING", (0, 0), (0, 0), 0.4 * cm),
        ("LEFTPADDING", (1, 0), (1, 0), 0.6 * cm), ("RIGHTPADDING", (1, 0), (1, 0), 0.6 * cm),
        ("TOPPADDING", (1, 0), (1, 0), 0.5 * cm), ("BOTTOMPADDING", (1, 0), (1, 0), 0.5 * cm),
    ]))

    story = _pdf_header_logo_flowables(cfg) + [header_table, Spacer(1, 0.3 * cm),
                                                HRFlowable(width="100%", thickness=1.2, color=PRIMARY, spaceAfter=6)]

    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"<b>Current Company:</b> {_esc(company_line)}", body))

    skills = candidate.get("skills") or []
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(", ".join(skills)), body))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(_esc(line_text), subhead))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def _render_pdf_timeline(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """Each role/section heading is marked with a colored bullet and a thin
    divider line right after it, giving a career-timeline feel while
    staying fully single-column and ATS-parseable -- no timeline graphic is
    tied to variable-length content, so it never risks a page-fit crash."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm,
                             topMargin=2 * cm, bottomMargin=2 * cm)
    ACCENT = colors.HexColor("#0d9488")
    DARK = colors.HexColor("#0f172a")
    GRAY = colors.HexColor("#64748b")

    h1 = ParagraphStyle("H1", fontSize=18, leading=22, textColor=DARK, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=4)
    sub = ParagraphStyle("Sub", fontSize=11, leading=14, textColor=ACCENT, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=10)
    h2 = ParagraphStyle("H2", fontSize=11, leading=14, textColor=ACCENT, fontName="Helvetica-Bold", spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("Body", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica")
    # The "timeline marker" -- a colored dot bullet baked into the text
    # (not reportlab's bulletText mechanism, since this needs a leading
    # dot character followed by normal left-aligned text, not a hanging
    # indent) plus a thin divider drawn right after via HRFlowable.
    timeline_head = ParagraphStyle("TimelineHead", fontSize=10, leading=15, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=2)
    bullet = ParagraphStyle("Bullet", fontSize=10, textColor=DARK, leading=15, fontName="Helvetica",
                             leftIndent=14, bulletIndent=0, spaceBefore=1, spaceAfter=1)
    small = ParagraphStyle("Small", fontSize=9, leading=11, textColor=GRAY, fontName="Helvetica")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    story = _pdf_header_logo_flowables(cfg) + [
        Paragraph(_esc(display_name), h1),
        Paragraph(_esc(candidate.get("current_designation") or ""), sub),
        HRFlowable(width="100%", thickness=1.2, color=ACCENT, spaceAfter=6),
    ]

    meta_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        meta_bits.append(f"<b>Mobile:</b> {_esc(candidate['phone'])}")
    if cfg["show_email"] and candidate.get("email"):
        meta_bits.append(f"<b>Email:</b> {_esc(candidate['email'])}")
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"<b>Location:</b> {_esc(candidate['location'])}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"<b>Experience:</b> {_esc(fmt_exp(candidate['total_exp_mo']))}")
    if meta_bits:
        story.append(Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(meta_bits), body))

    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"<b>Current Company:</b> {_esc(company_line)}", body))

    skills = candidate.get("skills") or []
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(", ".join(skills)), body))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(f'<font color="#0d9488">●</font> {_esc(line_text)}', timeline_head))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=4))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def _render_pdf_compact_grid(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """A real skills grid (a short, fixed-size Table, capped the same way
    the sidebar theme caps its own skill list) plus tighter body spacing --
    a denser, more information-forward layout."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                             topMargin=1.8 * cm, bottomMargin=1.8 * cm)
    PRIMARY = colors.HexColor("#1e40af")
    DARK = colors.HexColor("#0f172a")
    GRAY = colors.HexColor("#64748b")
    CHIP_BG = colors.HexColor("#f8fafc")

    h1 = ParagraphStyle("H1", fontSize=17, leading=20, textColor=DARK, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=3)
    sub = ParagraphStyle("Sub", fontSize=10.5, leading=13, textColor=PRIMARY, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=8)
    h2 = ParagraphStyle("H2", fontSize=10.5, leading=13, textColor=PRIMARY, fontName="Helvetica-Bold", spaceBefore=9, spaceAfter=4)
    # Real "compact" difference: leading=13 (vs 15 elsewhere) -- a genuinely
    # denser line-height, not just a label.
    body = ParagraphStyle("Body", fontSize=9.5, textColor=DARK, leading=13, fontName="Helvetica")
    subhead = ParagraphStyle("SubHead", fontSize=9.5, leading=13, textColor=DARK, fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=1)
    bullet = ParagraphStyle("Bullet", fontSize=9.5, textColor=DARK, leading=13, fontName="Helvetica",
                             leftIndent=12, bulletIndent=0, spaceBefore=0.5, spaceAfter=0.5)
    chip = ParagraphStyle("Chip", fontSize=8.5, leading=11, textColor=DARK, fontName="Helvetica", alignment=TA_CENTER)
    small = ParagraphStyle("Small", fontSize=8.5, leading=10, textColor=GRAY, fontName="Helvetica")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    story = _pdf_header_logo_flowables(cfg) + [Paragraph(_esc(display_name), h1)]
    if candidate.get("current_designation"):
        story.append(Paragraph(_esc(candidate["current_designation"]), sub))

    meta_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        meta_bits.append(_esc(candidate["phone"]))
    if cfg["show_email"] and candidate.get("email"):
        meta_bits.append(_esc(candidate["email"]))
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(_esc(candidate["location"]))
    if candidate.get("total_exp_mo"):
        meta_bits.append(_esc(fmt_exp(candidate["total_exp_mo"])) + " exp")
    if meta_bits:
        story.append(Paragraph(" &nbsp;•&nbsp; ".join(meta_bits), ParagraphStyle("MetaCenter", parent=body, alignment=TA_CENTER)))
    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"<b>Current Company:</b> {_esc(company_line)}", ParagraphStyle("CoCenter", parent=body, alignment=TA_CENTER)))
    story.append(HRFlowable(width="100%", thickness=1, color=PRIMARY, spaceBefore=6, spaceAfter=8))

    skills = (candidate.get("skills") or [])[:21]
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        cols = 3
        rows = [skills[i:i + cols] for i in range(0, len(skills), cols)]
        grid_data = [[Paragraph(_esc(s), chip) for s in row] + [""] * (cols - len(row)) for row in rows]
        skills_table = Table(grid_data, colWidths=[doc.width / cols] * cols)
        skills_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), CHIP_BG),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(skills_table)
        story.append(Spacer(1, 0.3 * cm))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(_esc(line_text), subhead))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def _render_pdf_elegant_serif(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    """Refined serif typography (reportlab's built-in Times family -- no
    external font embedding needed) with a double-rule header -- a
    premium, editorial look."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.6 * cm, rightMargin=2.6 * cm,
                             topMargin=2.2 * cm, bottomMargin=2.2 * cm)
    ACCENT = colors.HexColor("#7c2d12")
    DARK = colors.HexColor("#1c1917")
    GRAY = colors.HexColor("#78716c")

    h1 = ParagraphStyle("H1", fontSize=22, leading=26, textColor=DARK, fontName="Times-Bold", alignment=TA_CENTER, spaceAfter=3)
    sub = ParagraphStyle("Sub", fontSize=12, leading=15, textColor=ACCENT, fontName="Times-Italic", alignment=TA_CENTER, spaceAfter=8)
    h2 = ParagraphStyle("H2", fontSize=12, leading=15, textColor=ACCENT, fontName="Times-Bold", spaceBefore=13, spaceAfter=6)
    body = ParagraphStyle("Body", fontSize=10.5, textColor=DARK, leading=16, fontName="Times-Roman")
    subhead = ParagraphStyle("SubHead", fontSize=10.5, leading=16, textColor=DARK, fontName="Times-Bold", spaceBefore=8, spaceAfter=2)
    bullet = ParagraphStyle("Bullet", fontSize=10.5, textColor=DARK, leading=16, fontName="Times-Roman",
                             leftIndent=15, bulletIndent=0, spaceBefore=1, spaceAfter=1)
    small = ParagraphStyle("Small", fontSize=9, leading=11, textColor=GRAY, fontName="Times-Italic")

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    story = _pdf_header_logo_flowables(cfg) + [
        Paragraph(_esc(display_name), h1),
        Paragraph(_esc(candidate.get("current_designation") or ""), sub),
        HRFlowable(width="100%", thickness=1.4, color=ACCENT, spaceAfter=1.5),
        HRFlowable(width="100%", thickness=0.5, color=ACCENT, spaceAfter=8),
    ]

    meta_bits = []
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"<b>Location:</b> {_esc(candidate['location'])}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"<b>Experience:</b> {_esc(fmt_exp(candidate['total_exp_mo']))}")
    if meta_bits:
        story.append(Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(meta_bits), body))
    contact_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        contact_bits.append(f"<b>Mobile:</b> {_esc(candidate['phone'])}")
    if cfg["show_email"] and candidate.get("email"):
        contact_bits.append(f"<b>Email:</b> {_esc(candidate['email'])}")
    if contact_bits:
        story.append(Paragraph(" &nbsp;&nbsp;|&nbsp;&nbsp; ".join(contact_bits), body))

    company_line = _company_line(candidate, cfg)
    if company_line:
        story.append(Paragraph(f"<b>Current Company:</b> {_esc(company_line)}", body))

    skills = candidate.get("skills") or []
    if skills:
        story.append(Paragraph("KEY SKILLS", h2))
        story.append(Paragraph(_esc(", ".join(skills)), body))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        story.append(Paragraph(heading, h2))
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                story.append(Paragraph(_esc(line_text), bullet, bulletText='•'))
            elif kind == 'subhead':
                story.append(Paragraph(_esc(line_text), subhead))
            else:
                story.append(Paragraph(_esc(line_text), body))

    client_line = _client_line(client_name, cfg)
    story.extend(_pdf_footer_flowables(cfg, client_line, small))
    doc.build(story)
    return buf.getvalue()


def _render_docx_executive_header(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    BAND_TEXT = RGBColor(0xff, 0xff, 0xff)
    PRIMARY = RGBColor(0x1e, 0x40, 0xaf)
    DARK = RGBColor(0x0f, 0x17, 0x2a)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    _docx_header_logo(doc, cfg)

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    band = doc.add_table(rows=1, cols=1)
    cell = band.rows[0].cells[0]
    _shade_cell(cell, "1E3A5F")
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BAND_TEXT
    if candidate.get("current_designation"):
        p2 = cell.add_paragraph()
        r2 = p2.add_run(candidate["current_designation"])
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0xcb, 0xd5, 0xe1)
    doc.add_paragraph()

    meta_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        meta_bits.append(f"Mobile: {candidate['phone']}")
    if cfg["show_email"] and candidate.get("email"):
        meta_bits.append(f"Email: {candidate['email']}")
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"Location: {candidate['location']}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"Experience: {fmt_exp(candidate['total_exp_mo'])}")
    if meta_bits:
        doc.add_paragraph(" | ".join(meta_bits))

    company_line = _company_line(candidate, cfg)
    if company_line:
        doc.add_paragraph(f"Current Company: {company_line}")

    skills = candidate.get("skills") or []
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.color.rgb = PRIMARY
        doc.add_paragraph(", ".join(skills))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = PRIMARY
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                doc.add_paragraph(line_text, style='List Bullet')
                continue
            pp = doc.add_paragraph()
            r = pp.add_run(line_text)
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = DARK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_two_tone_header(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    PRIMARY = RGBColor(0x1e, 0x40, 0xaf)
    DARK = RGBColor(0x0f, 0x17, 0x2a)

    doc = Document()
    _docx_header_logo(doc, cfg)

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(11.5)
    right.width = Cm(7)
    _shade_cell(right, "EFF6FF")

    left.paragraphs[0].text = ""
    p = left.paragraphs[0]
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK
    if candidate.get("current_designation"):
        p2 = left.add_paragraph()
        r2 = p2.add_run(candidate["current_designation"])
        r2.bold = True
        r2.font.size = Pt(11.5)
        r2.font.color.rgb = PRIMARY

    right.paragraphs[0].text = ""
    right_lines = []
    if cfg["show_mobile"] and candidate.get("phone"):
        right_lines.append(candidate["phone"])
    if cfg["show_email"] and candidate.get("email"):
        right_lines.append(candidate["email"])
    if cfg["show_location"] and candidate.get("location"):
        right_lines.append(candidate["location"])
    if candidate.get("total_exp_mo"):
        right_lines.append(f"{fmt_exp(candidate['total_exp_mo'])} experience")
    if right_lines:
        rp = right.paragraphs[0]
        rr = rp.add_run(right_lines[0])
        rr.font.size = Pt(9)
        for extra in right_lines[1:]:
            rp2 = right.add_paragraph()
            rr2 = rp2.add_run(extra)
            rr2.font.size = Pt(9)

    doc.add_paragraph()
    company_line = _company_line(candidate, cfg)
    if company_line:
        doc.add_paragraph(f"Current Company: {company_line}")

    skills = candidate.get("skills") or []
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.color.rgb = PRIMARY
        doc.add_paragraph(", ".join(skills))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = PRIMARY
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                doc.add_paragraph(line_text, style='List Bullet')
                continue
            pp = doc.add_paragraph()
            r = pp.add_run(line_text)
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = DARK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_timeline(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ACCENT = RGBColor(0x0d, 0x94, 0x88)
    DARK = RGBColor(0x0f, 0x17, 0x2a)

    doc = Document()
    _docx_header_logo(doc, cfg)

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = DARK
    if candidate.get("current_designation"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(candidate["current_designation"])
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = ACCENT

    meta_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        meta_bits.append(f"Mobile: {candidate['phone']}")
    if cfg["show_email"] and candidate.get("email"):
        meta_bits.append(f"Email: {candidate['email']}")
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"Location: {candidate['location']}")
    if meta_bits:
        doc.add_paragraph(" | ".join(meta_bits))

    company_line = _company_line(candidate, cfg)
    if company_line:
        doc.add_paragraph(f"Current Company: {company_line}")

    skills = candidate.get("skills") or []
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.color.rgb = ACCENT
        doc.add_paragraph(", ".join(skills))

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = ACCENT
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                doc.add_paragraph(line_text, style='List Bullet')
                continue
            pp = doc.add_paragraph()
            if kind == 'subhead':
                r1 = pp.add_run("● ")
                r1.font.color.rgb = ACCENT
                r2 = pp.add_run(line_text)
                r2.bold = True
                r2.font.color.rgb = DARK
            else:
                pp.add_run(line_text)

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_compact_grid(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    PRIMARY = RGBColor(0x1e, 0x40, 0xaf)
    DARK = RGBColor(0x0f, 0x17, 0x2a)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    _docx_header_logo(doc, cfg)

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK
    if candidate.get("current_designation"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(candidate["current_designation"])
        r2.bold = True
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = PRIMARY

    meta_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        meta_bits.append(candidate["phone"])
    if cfg["show_email"] and candidate.get("email"):
        meta_bits.append(candidate["email"])
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(candidate["location"])
    if meta_bits:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.add_run(" • ".join(meta_bits)).font.size = Pt(9.5)

    skills = (candidate.get("skills") or [])[:21]
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.color.rgb = PRIMARY
        cols = 3
        rows = [skills[i:i + cols] for i in range(0, len(skills), cols)]
        table = doc.add_table(rows=len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci in range(cols):
                cell = table.rows[ri].cells[ci]
                cell.text = row[ci] if ci < len(row) else ""
                for pp in cell.paragraphs:
                    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for rr in pp.runs:
                        rr.font.size = Pt(9)

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.color.rgb = PRIMARY
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                bp = doc.add_paragraph(line_text, style='List Bullet')
                for rr in bp.runs:
                    rr.font.size = Pt(9.5)
                continue
            pp = doc.add_paragraph()
            r = pp.add_run(line_text)
            r.font.size = Pt(9.5)
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = DARK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_elegant_serif(candidate: dict, cfg: dict, client_name: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ACCENT = RGBColor(0x7c, 0x2d, 0x12)
    DARK = RGBColor(0x1c, 0x19, 0x17)
    SERIF = "Times New Roman"

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

    _docx_header_logo(doc, cfg)

    display_name = mask_name(candidate.get("full_name") or "") if cfg["name_format"] == "masked" else (candidate.get("full_name") or "Candidate")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = SERIF
    r.font.color.rgb = DARK
    if candidate.get("current_designation"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(candidate["current_designation"])
        r2.italic = True
        r2.font.size = Pt(12)
        r2.font.name = SERIF
        r2.font.color.rgb = ACCENT

    meta_bits = []
    if cfg["show_location"] and candidate.get("location"):
        meta_bits.append(f"Location: {candidate['location']}")
    if candidate.get("total_exp_mo"):
        meta_bits.append(f"Experience: {fmt_exp(candidate['total_exp_mo'])}")
    if meta_bits:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = mp.add_run(" | ".join(meta_bits))
        rr.font.name = SERIF

    contact_bits = []
    if cfg["show_mobile"] and candidate.get("phone"):
        contact_bits.append(f"Mobile: {candidate['phone']}")
    if cfg["show_email"] and candidate.get("email"):
        contact_bits.append(f"Email: {candidate['email']}")
    if contact_bits:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = cp.add_run(" | ".join(contact_bits))
        rr.font.name = SERIF

    company_line = _company_line(candidate, cfg)
    if company_line:
        p3 = doc.add_paragraph()
        r3 = p3.add_run(f"Current Company: {company_line}")
        r3.font.name = SERIF

    skills = candidate.get("skills") or []
    if skills:
        h = doc.add_paragraph()
        r = h.add_run("KEY SKILLS")
        r.bold = True
        r.font.name = SERIF
        r.font.color.rgb = ACCENT
        sp = doc.add_paragraph()
        sr = sp.add_run(", ".join(skills))
        sr.font.name = SERIF

    heading, text = _resolve_body_text(candidate, cfg)
    if text.strip():
        h = doc.add_paragraph()
        r = h.add_run(heading)
        r.bold = True
        r.font.name = SERIF
        r.font.color.rgb = ACCENT
        for line_text, kind in _classify_lines(text):
            if kind == 'bullet':
                bp = doc.add_paragraph(line_text, style='List Bullet')
                for rr in bp.runs:
                    rr.font.name = SERIF
                continue
            pp = doc.add_paragraph()
            r = pp.add_run(line_text)
            r.font.name = SERIF
            if kind == 'subhead':
                r.bold = True
                r.font.color.rgb = DARK

    client_line = _client_line(client_name, cfg)
    _docx_footer(doc, cfg, client_line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Populated here, after every _render_pdf_*/_render_docx_* function above is
# defined -- referenced by name (a dict literal, not a decorator), so the
# order this module is imported in doesn't matter, only that this line runs
# after the function defs above it, which it does by construction.
_PDF_RENDERERS.update({
    "classic": _render_pdf_classic,
    "modern_sidebar": _render_pdf_sidebar,
    "minimal_ats": _render_pdf_minimal,
    "executive_header": _render_pdf_executive_header,
    "two_tone_header": _render_pdf_two_tone_header,
    "timeline": _render_pdf_timeline,
    "compact_grid": _render_pdf_compact_grid,
    "elegant_serif": _render_pdf_elegant_serif,
})
_DOCX_RENDERERS.update({
    "classic": _render_docx_classic,
    "modern_sidebar": _render_docx_sidebar,
    "minimal_ats": _render_docx_minimal,
    "executive_header": _render_docx_executive_header,
    "two_tone_header": _render_docx_two_tone_header,
    "timeline": _render_docx_timeline,
    "compact_grid": _render_docx_compact_grid,
    "elegant_serif": _render_docx_elegant_serif,
})
