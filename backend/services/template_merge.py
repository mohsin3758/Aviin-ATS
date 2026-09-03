"""Fills an uploaded .xlsx/.docx tracking-sheet template with real
candidate/requisition data via {{token}} placeholders.

Scope, stated plainly rather than silently under-delivered: only .xlsx and
.docx support real merge-fill (both are structured, text-addressable
formats openpyxl/python-docx can read a placeholder out of and write a real
value back into). An uploaded .pdf CANNOT be reliably merge-filled — a PDF
is a flattened, rendered page with no addressable "this text run is a
field" concept the way a Word/Excel document has — so a .pdf template is
used as a static reference attachment only (sent alongside, never as the
live data carrier). This mirrors this codebase's own precedent of never
claiming a capability that doesn't actually work (see resume_formatting.py's
documented scope boundaries).

Token format: {{key}}, where key is any tracking_sheet_templates column key
(a COLUMN_REGISTRY key or a KAE-typed custom column key). A template can use
any subset of tokens, in any order, anywhere in the document.

Two real, distinct merge strategies:
  - Single-row templates (no repeated {{sl_no}}-style row block detected):
    every token anywhere in the document is replaced with the MOST RECENT
    row's value — a cover-letter-style document, not a cumulative sheet.
  - Row-block templates (.xlsx: a single spreadsheet row containing tokens;
    .docx: a single table row containing tokens): that row is genuinely
    duplicated once per real row passed in and each copy's tokens are
    replaced with that row's own values. A caller sending a single row
    (the real, current default for every KAE-submission send since
    2026-09-03 — one candidate's own tracking row per email, never a
    growing history of prior candidates) fills exactly one row here too.
"""
import io
import re
from copy import deepcopy

TOKEN_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def _substitute(text: str, values: dict) -> str:
    def repl(m):
        key = m.group(1)
        v = values.get(key)
        return "" if v is None else str(v)
    return TOKEN_RE.sub(repl, text)


def _row_has_token(cell_texts: list[str]) -> bool:
    return any(TOKEN_RE.search(t or "") for t in cell_texts)


# ─────────────────────────── XLSX ───────────────────────────

def fill_xlsx_template(file_bytes: bytes, rows: list[dict]) -> bytes:
    from openpyxl import load_workbook
    from copy import copy as style_copy

    wb = load_workbook(io.BytesIO(file_bytes))
    ws = wb.active

    # Find the first row containing at least one {{token}} cell — that's the
    # real template/data row. Everything above it (title, headers) is left
    # untouched.
    template_row_idx = None
    for r in range(1, ws.max_row + 1):
        texts = [str(ws.cell(row=r, column=c).value or "") for c in range(1, ws.max_column + 1)]
        if _row_has_token(texts):
            template_row_idx = r
            break

    if template_row_idx is None or not rows:
        # No token row found (or nothing to fill) — return the file
        # untouched rather than guessing at a structure that isn't there.
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    max_col = ws.max_column
    template_cells = [ws.cell(row=template_row_idx, column=c) for c in range(1, max_col + 1)]
    template_texts = [str(c.value or "") for c in template_cells]
    template_styles = [style_copy(c._style) for c in template_cells]

    # Fill the template row itself with the first data row's values, then
    # insert one new, style-matching row per additional data row directly
    # below it (openpyxl has no "duplicate row" primitive — insert a blank
    # row and copy each cell's style + substituted text across by hand).
    for c_idx, cell in enumerate(template_cells, start=1):
        ws.cell(row=template_row_idx, column=c_idx).value = _substitute(template_texts[c_idx - 1], rows[0])

    for offset, row_values in enumerate(rows[1:], start=1):
        insert_at = template_row_idx + offset
        ws.insert_rows(insert_at)
        for c_idx in range(1, max_col + 1):
            new_cell = ws.cell(row=insert_at, column=c_idx)
            new_cell._style = style_copy(template_styles[c_idx - 1])
            new_cell.value = _substitute(template_texts[c_idx - 1], row_values)
        ws.row_dimensions[insert_at].height = ws.row_dimensions[template_row_idx].height

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ─────────────────────────── DOCX ───────────────────────────

def fill_docx_template(file_bytes: bytes, rows: list[dict]) -> bytes:
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    latest = rows[-1] if rows else {}

    # Plain-paragraph tokens anywhere in the body (outside any table) always
    # use the most recent row's values — a cover note, not a repeating list.
    for p in doc.paragraphs:
        if TOKEN_RE.search(p.text):
            _replace_paragraph_tokens(p, latest)

    for table in doc.tables:
        token_row = None
        for tr in table.rows:
            cell_texts = [c.text for c in tr.cells]
            if _row_has_token(cell_texts):
                token_row = tr
                break
        if token_row is None:
            continue

        # First data row fills the template row itself; every additional row
        # is a real, deep-copied duplicate of the template row's XML,
        # appended to the table, with its own tokens substituted — a genuine
        # cumulative table, not just a single filled-in row.
        cell_templates = [c.text for c in token_row.cells]
        for c, tmpl_text in zip(token_row.cells, cell_templates):
            _set_cell_text(c, _substitute(tmpl_text, rows[0] if rows else {}))

        for row_values in rows[1:]:
            new_tr_elem = deepcopy(token_row._tr)
            token_row._tr.addnext(new_tr_elem)
            from docx.table import _Row
            new_row = _Row(new_tr_elem, table)
            for c, tmpl_text in zip(new_row.cells, cell_templates):
                _set_cell_text(c, _substitute(tmpl_text, row_values))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_paragraph_tokens(paragraph, values: dict):
    """Token text can be split across multiple runs by Word's own editing
    history — collapse the paragraph's runs into one before substituting,
    same real-world caveat every docx mail-merge tool has to handle."""
    full = paragraph.text
    if not TOKEN_RE.search(full):
        return
    new_text = _substitute(full, values)
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def _set_cell_text(cell, text: str):
    """Same run-collapse approach as _replace_paragraph_tokens, applied to
    every paragraph inside the cell (a table cell can itself contain
    multiple paragraphs, though a tracking-sheet cell almost never does)."""
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p = cell.paragraphs[0]
    for run in p.runs[1:]:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p.text = ""
