"""NDA document generation + e-signature (Stage-Workflow Automation Phase 1).

Mirrors the offer-letter pattern in offers.py, which is already live in
production: reportlab PDF, a public no-auth signing link backed by
SECURITY DEFINER SQL functions (sql/12_nda_esign.sql) so RLS stays intact
without exposing tenant context to anonymous requests.

HARD RULE #12: a consent_records row is written before the NDA is shared
with the candidate. HARD RULE #5/#6: the auto-advance stage change writes
event_outbox in the same transaction with a dedup_key. Not HITL-gated —
"candidate rejected / offer issued / recruiter reassigned" are the only
actions HARD RULE #10 requires a pause for; NDA auto-advance is not one of
them, so it proceeds automatically as requested.
"""

import os
import secrets
import smtplib
from datetime import date
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from io import BytesIO
from pathlib import Path
from typing import Optional

import asyncpg
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

import db
import events
from deps import Actor, get_actor

router = APIRouter(prefix="/applications", tags=["nda"])
nda_router = APIRouter(prefix="/nda", tags=["nda"])
nda_sign_public = APIRouter(prefix="/nda-sign", tags=["nda-sign"])
doc_templates_router = APIRouter(prefix="/settings/document-templates", tags=["document-templates"])

NDA_FIELDS = """id, tenant_id, application_id, candidate_id, draft_text, final_text,
                status, sign_method, signatory_name, sent_at, signed_at, created_at,
                manual_file_path, attachment_source, attached_file_name"""

UPLOAD_DIR = Path("/app/uploads/nda")
TEMPLATE_UPLOAD_DIR = Path("/app/uploads/document_templates")
ALLOWED_TEMPLATE_EXTS = {".pdf", ".doc", ".docx"}


def _default_nda_text(candidate_name: str, job_title: str, company_name: str) -> str:
    today = date.today().strftime("%d %B %Y")
    return (
        f"This Non-Disclosure and Pre-Contract Agreement (\"Agreement\") is entered into "
        f"between {company_name} (\"Company\") and {candidate_name} (\"Candidate\") as of {today}.\n\n"
        f"1. CONFIDENTIALITY: The Candidate agrees to keep confidential all information shared "
        f"regarding the role of {job_title}, including client identity, compensation details, "
        f"and process specifics, and not disclose it to any third party.\n\n"
        f"2. NON-CIRCUMVENTION: The Candidate agrees not to approach the client directly or "
        f"through any other channel outside of {company_name} for this or related opportunities "
        f"during the hiring process.\n\n"
        f"3. ACCURACY: The Candidate confirms that all information, documents, and credentials "
        f"shared with {company_name} are true and accurate to the best of their knowledge.\n\n"
        f"4. DURATION: This Agreement remains in effect through the duration of the recruitment "
        f"process and for twelve (12) months thereafter.\n\n"
        f"By signing below, the Candidate acknowledges having read, understood, and agreed to "
        f"the terms of this Agreement."
    )


async def _nda_context(conn, application_id: str) -> dict:
    row = await conn.fetchrow(
        """SELECT c.full_name AS candidate_name, c.email AS candidate_email,
                  r.title AS job_title, t.name AS company_name
           FROM applications a
           JOIN candidates c ON c.id = a.candidate_id
           JOIN requisitions r ON r.id = a.requisition_id
           JOIN tenants t ON t.id = a.tenant_id
           WHERE a.id = $1""",
        application_id,
    )
    return dict(row) if row else {}


async def _get_or_create_nda(conn, tenant_id: str, application_id: str):
    row = await conn.fetchrow(f"SELECT {NDA_FIELDS} FROM nda_documents WHERE application_id=$1", application_id)
    if row:
        return row
    ctx = await _nda_context(conn, application_id)
    if not ctx:
        raise HTTPException(404, "Application not found")
    app_row = await conn.fetchrow("SELECT candidate_id FROM applications WHERE id=$1", application_id)
    draft = _default_nda_text(
        ctx.get("candidate_name") or "Candidate",
        ctx.get("job_title") or "the role",
        ctx.get("company_name") or "AVIIN Jobs Services",
    )
    return await conn.fetchrow(
        f"""INSERT INTO nda_documents (tenant_id, application_id, candidate_id, draft_text, status)
            VALUES ($1,$2,$3,$4,'draft') RETURNING {NDA_FIELDS}""",
        tenant_id, application_id, app_row["candidate_id"], draft,
    )


# ─── Document Templates (upload your own NDA / Contract file, PDF or Word) ───
# Reusable per tenant: upload once, replace or remove any time, used as the
# attachment when sending a candidate's NDA instead of the auto-generated PDF.

MIME_BY_EXT = {
    ".pdf": "application/pdf",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@doc_templates_router.get("")
async def list_document_templates(actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        rows = await conn.fetch(
            "SELECT doc_type, file_name, mime_type, uploaded_at FROM document_templates WHERE tenant_id=$1",
            actor.tenant_id)
    by_type = {r["doc_type"]: dict(r) for r in rows}
    return {"nda": by_type.get("nda"), "contract": by_type.get("contract")}


@doc_templates_router.post("/{doc_type}")
async def upload_document_template(doc_type: str, file: UploadFile = File(...), actor: Actor = Depends(get_actor)):
    if doc_type not in ("nda", "contract"):
        raise HTTPException(400, "doc_type must be 'nda' or 'contract'")
    ext = "." + (file.filename or "").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_TEMPLATE_EXTS:
        raise HTTPException(400, f"Unsupported file type: {ext}. Use PDF or Word (.pdf, .doc, .docx)")
    file_bytes = await file.read()

    folder = TEMPLATE_UPLOAD_DIR / actor.tenant_id
    folder.mkdir(parents=True, exist_ok=True)
    rel_path = f"/uploads/document_templates/{actor.tenant_id}/{doc_type}{ext}"
    (folder / f"{doc_type}{ext}").write_bytes(file_bytes)

    async with db.tenant_conn(actor.tenant_id) as conn:
        row = await conn.fetchrow(
            """INSERT INTO document_templates (tenant_id, doc_type, file_path, file_name, mime_type, uploaded_by)
               VALUES ($1,$2,$3,$4,$5,$6)
               ON CONFLICT (tenant_id, doc_type) DO UPDATE SET
                 file_path=EXCLUDED.file_path, file_name=EXCLUDED.file_name,
                 mime_type=EXCLUDED.mime_type, uploaded_by=EXCLUDED.uploaded_by, uploaded_at=now()
               RETURNING doc_type, file_name, mime_type, uploaded_at""",
            actor.tenant_id, doc_type, rel_path, file.filename or f"{doc_type}{ext}",
            MIME_BY_EXT.get(ext, "application/octet-stream"), actor.user_id,
        )
    return dict(row)


@doc_templates_router.delete("/{doc_type}")
async def remove_document_template(doc_type: str, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        row = await conn.fetchrow(
            "DELETE FROM document_templates WHERE tenant_id=$1 AND doc_type=$2 RETURNING file_path",
            actor.tenant_id, doc_type)
    if row and row["file_path"]:
        try:
            (Path("/app") / row["file_path"].lstrip("/")).unlink(missing_ok=True)
        except Exception:
            pass
    return {"removed": bool(row)}


@doc_templates_router.get("/{doc_type}/download")
async def download_document_template(doc_type: str, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        row = await conn.fetchrow(
            "SELECT file_path, file_name FROM document_templates WHERE tenant_id=$1 AND doc_type=$2",
            actor.tenant_id, doc_type)
    if not row:
        raise HTTPException(404, "No template uploaded for this type")
    abs_path = Path("/app") / row["file_path"].lstrip("/")
    if not abs_path.exists():
        raise HTTPException(404, "File missing from disk")
    return FileResponse(str(abs_path), filename=row["file_name"])


# ─── PDF / DOCX generation (reportlab / python-docx, zero-token — static template) ──

def _build_nda_pdf(text: str, candidate_name: str, company_name: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2.5*cm, rightMargin=2.5*cm,
                             topMargin=2*cm, bottomMargin=2*cm)
    PRIMARY = colors.HexColor('#1e40af')
    DARK = colors.HexColor('#0f172a')
    GRAY = colors.HexColor('#64748b')
    h1 = ParagraphStyle('H1', fontSize=18, textColor=PRIMARY, spaceAfter=4,
                         fontName='Helvetica-Bold', alignment=TA_CENTER)
    title_style = ParagraphStyle('Title', fontSize=13, textColor=DARK, spaceAfter=14,
                                  fontName='Helvetica-Bold', alignment=TA_CENTER)
    body = ParagraphStyle('Body', fontSize=10, textColor=DARK, spaceAfter=8,
                           leading=16, fontName='Helvetica', alignment=TA_JUSTIFY)
    small = ParagraphStyle('Small', fontSize=9, textColor=GRAY, spaceAfter=4, fontName='Helvetica')

    story = [
        Paragraph(company_name, h1),
        HRFlowable(width='100%', thickness=2, color=PRIMARY, spaceAfter=12),
        Paragraph('NON-DISCLOSURE / PRE-CONTRACT AGREEMENT', title_style),
    ]
    for para in text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip().replace("\n", "<br/>"), body))
            story.append(Spacer(1, 0.15*cm))
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width='100%', thickness=1, color=GRAY, spaceAfter=10))
    story.append(Paragraph('Signature: _________________________________   Date: _________________', small))
    story.append(Paragraph(f'Candidate: {candidate_name}', small))

    doc.build(story)
    return buf.getvalue()


def _build_nda_docx(text: str, candidate_name: str, company_name: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = Document()
    h = d.add_heading(company_name, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = d.add_heading('Non-Disclosure / Pre-Contract Agreement', level=2)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for para in text.split("\n\n"):
        if para.strip():
            p = d.add_paragraph(para.strip())
            p.paragraph_format.space_after = Pt(10)
    d.add_paragraph("")
    d.add_paragraph("Signature: _________________________________   Date: _________________")
    d.add_paragraph(f"Candidate: {candidate_name}")

    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


# ─── Draft CRUD ─────────────────────────────────────────────────────────────

@router.get("/{application_id}/nda")
async def get_nda(application_id: str, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        row = await _get_or_create_nda(conn, actor.tenant_id, application_id)
    return dict(row)


class NdaTextUpdate(BaseModel):
    draft_text: str


@router.put("/{application_id}/nda")
async def save_nda_draft(application_id: str, body: NdaTextUpdate, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        await _get_or_create_nda(conn, actor.tenant_id, application_id)
        row = await conn.fetchrow(
            f"UPDATE nda_documents SET draft_text=$1 WHERE application_id=$2 RETURNING {NDA_FIELDS}",
            body.draft_text, application_id,
        )
    return dict(row)


@router.get("/{application_id}/nda/pdf")
async def download_nda_pdf(application_id: str, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        nda = await _get_or_create_nda(conn, actor.tenant_id, application_id)
        ctx = await _nda_context(conn, application_id)
    text = nda["final_text"] or nda["draft_text"]
    pdf_bytes = _build_nda_pdf(text, ctx.get("candidate_name", "Candidate"), ctx.get("company_name", "AVIIN Jobs Services"))
    fname = f"nda_{application_id[:8]}.pdf"
    return StreamingResponse(BytesIO(pdf_bytes), media_type='application/pdf',
                              headers={'Content-Disposition': f'attachment; filename="{fname}"'})


@router.get("/{application_id}/nda/docx")
async def download_nda_docx(application_id: str, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        nda = await _get_or_create_nda(conn, actor.tenant_id, application_id)
        ctx = await _nda_context(conn, application_id)
    text = nda["final_text"] or nda["draft_text"]
    docx_bytes = _build_nda_docx(text, ctx.get("candidate_name", "Candidate"), ctx.get("company_name", "AVIIN Jobs Services"))
    fname = f"nda_{application_id[:8]}.docx"
    return StreamingResponse(
        BytesIO(docx_bytes),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'})


# ─── Send for signature ─────────────────────────────────────────────────────

async def _send_email_with_pdf(tenant_id: str, to_email: str, to_name: str, subject: str, body_text: str,
                                pdf_bytes: Optional[bytes] = None, pdf_filename: str = "document.pdf",
                                attachment_mime: str = "application/pdf"):
    """Best-effort SMTP send using this tenant's active email_settings row.

    Mirrors the working pattern in applications.py::_notify_stage_change_bg
    (smtp_password / smtp_from — NOT the smtp_pass / from_email columns used
    in offers.py::send_offer_letter, which don't exist and silently no-op).
    """
    try:
        db_url = os.environ.get("DATABASE_URL", "postgresql://app_user:apppw@db:5432/ats")
        conn = await asyncpg.connect(db_url)
        try:
            cfg = await conn.fetchrow(
                "SELECT smtp_host,smtp_port,smtp_user,smtp_password,smtp_from,smtp_from_name,smtp_tls "
                "FROM email_settings WHERE tenant_id=$1 AND is_active=TRUE LIMIT 1", tenant_id)
        finally:
            await conn.close()
        if not cfg or not cfg["smtp_host"]:
            print("NDA email: no active SMTP config found")
            return False
        msg = MIMEMultipart()
        msg["Subject"] = subject
        msg["From"] = f'{cfg["smtp_from_name"] or "AVIIN ATS"} <{cfg["smtp_from"] or cfg["smtp_user"]}>'
        msg["To"] = to_email
        msg.attach(MIMEText(body_text, "plain"))
        if pdf_bytes:
            maintype, _, subtype = attachment_mime.partition('/')
            part = MIMEBase(maintype or 'application', subtype or 'octet-stream')
            part.set_payload(pdf_bytes)
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename="{pdf_filename}"')
            msg.attach(part)
        port = cfg["smtp_port"] or 587
        with smtplib.SMTP(cfg["smtp_host"], port, timeout=10) as s:
            s.ehlo()
            if cfg["smtp_tls"] and port == 587:
                s.starttls()
                s.ehlo()
            if cfg["smtp_user"]:
                s.login(cfg["smtp_user"], cfg["smtp_password"] or "")
            s.sendmail(cfg["smtp_from"] or cfg["smtp_user"], [to_email], msg.as_string())
        return True
    except Exception as exc:
        print(f"NDA email failed to {to_email}: {exc}")
        return False


class NdaSendRequest(BaseModel):
    sign_method: str = "type_name"        # 'type_name' | 'otp'
    attachment: str = "generated"          # 'generated' | 'nda_template' | 'contract_template'


@router.post("/{application_id}/nda/send")
async def send_nda(application_id: str, body: NdaSendRequest, actor: Actor = Depends(get_actor)):
    if body.sign_method not in ("type_name", "otp"):
        raise HTTPException(400, "sign_method must be 'type_name' or 'otp'")
    if body.attachment not in ("generated", "nda_template", "contract_template"):
        raise HTTPException(400, "attachment must be 'generated', 'nda_template', or 'contract_template'")

    async with db.tenant_conn(actor.tenant_id) as conn:
        nda = await _get_or_create_nda(conn, actor.tenant_id, application_id)
        ctx = await _nda_context(conn, application_id)
        if not ctx.get("candidate_email"):
            raise HTTPException(400, "Candidate has no email address")

        attach_bytes: Optional[bytes] = None
        attach_filename = "nda_agreement.pdf"
        attach_mime = "application/pdf"
        attached_path: Optional[str] = None
        attached_name: Optional[str] = None

        if body.attachment != "generated":
            tmpl_type = "nda" if body.attachment == "nda_template" else "contract"
            tmpl = await conn.fetchrow(
                "SELECT file_path, file_name, mime_type FROM document_templates WHERE tenant_id=$1 AND doc_type=$2",
                actor.tenant_id, tmpl_type)
            if not tmpl:
                raise HTTPException(400, f"No {tmpl_type} template uploaded — upload one first or send the auto-generated document")
            abs_path = Path("/app") / tmpl["file_path"].lstrip("/")
            if not abs_path.exists():
                raise HTTPException(400, "Template file missing from disk — re-upload it")
            attach_bytes = abs_path.read_bytes()
            attach_filename = tmpl["file_name"]
            attach_mime = tmpl["mime_type"]
            attached_path = tmpl["file_path"]
            attached_name = tmpl["file_name"]

        token = secrets.token_urlsafe(32)
        final_text = nda["final_text"] or nda["draft_text"]
        nda = await conn.fetchrow(
            f"""UPDATE nda_documents
                SET final_text=$1, status='sent', sign_method=$2, signing_token=$3, sent_at=now(),
                    attachment_source=$4, attached_file_path=$5, attached_file_name=$6
                WHERE application_id=$7 RETURNING {NDA_FIELDS}""",
            final_text, body.sign_method, token, body.attachment, attached_path, attached_name, application_id,
        )

        # HARD RULE #12: consent record before sharing/processing candidate PII.
        await conn.execute(
            "INSERT INTO consent_records (tenant_id,candidate_id,data_category,channel,consent_given,consent_text) "
            "VALUES ($1,$2,'nda_agreement','email',TRUE,$3)",
            actor.tenant_id, nda["candidate_id"],
            f"NDA sent for e-signature via {body.sign_method} method on {date.today().isoformat()}",
        )

    base = os.environ.get("NEXT_PUBLIC_APP_URL", "https://ats.aviinjobs.com")
    sign_url = f"{base}/sign-nda/{token}"
    if attach_bytes is None:
        attach_bytes = _build_nda_pdf(final_text, ctx["candidate_name"], ctx.get("company_name", "AVIIN Jobs Services"))
    body_text = (
        f'Dear {ctx["candidate_name"]},\n\n'
        f'As part of our recruitment process for {ctx.get("job_title", "this role")}, please review and '
        f'sign the attached NDA / Pre-Contract Agreement.\n\n'
        f'Sign online here: {sign_url}\n\n'
        f'Best regards,\n{ctx.get("company_name", "AVIIN Jobs Services")}'
    )
    import asyncio
    asyncio.create_task(_send_email_with_pdf(
        actor.tenant_id, ctx["candidate_email"], ctx["candidate_name"],
        f'{ctx.get("company_name", "AVIIN Jobs Services")} - NDA / Pre-Contract Agreement',
        body_text, attach_bytes, attach_filename, attach_mime,
    ))
    return {"sent": True, "sign_url": sign_url, "recipient": ctx["candidate_email"]}


# ─── Manual sign (recruiter uploads a scanned signed copy) ──────────────────

ALLOWED_MANUAL_EXTS = {".pdf", ".jpg", ".jpeg", ".png", ".doc", ".docx"}


@router.post("/{application_id}/nda/manual-sign")
async def manual_sign_nda(application_id: str, file: UploadFile = File(...),
                           signatory_name: Optional[str] = None,
                           actor: Actor = Depends(get_actor)):
    ext = "." + (file.filename or "signed").rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_MANUAL_EXTS:
        raise HTTPException(400, f"Unsupported file type: {ext}")
    file_bytes = await file.read()

    async with db.tenant_conn(actor.tenant_id) as conn:
        nda = await _get_or_create_nda(conn, actor.tenant_id, application_id)

        folder = UPLOAD_DIR / actor.tenant_id
        folder.mkdir(parents=True, exist_ok=True)
        rel_path = f"/uploads/nda/{actor.tenant_id}/{nda['id']}{ext}"
        (UPLOAD_DIR / actor.tenant_id / f"{nda['id']}{ext}").write_bytes(file_bytes)

        await conn.execute(
            """UPDATE nda_documents
               SET status='manually_signed', sign_method='manual', signatory_name=$1,
                   manual_file_path=$2, uploaded_by=$3, signed_at=now()
               WHERE application_id=$4""",
            signatory_name or "Signed copy uploaded by recruiter", rel_path, actor.user_id, application_id,
        )

    await _on_nda_signed(actor.tenant_id, application_id)
    return {"signed": True, "method": "manual"}


@nda_router.get("")
async def list_nda_documents(status: Optional[str] = None, actor: Actor = Depends(get_actor)):
    """All NDA documents for this tenant — powers the /nda-documents list page."""
    async with db.tenant_conn(actor.tenant_id) as conn:
        rows = await conn.fetch(
            """SELECT nd.id, nd.application_id, nd.status, nd.sign_method, nd.signatory_name,
                      nd.sent_at, nd.signed_at, nd.created_at, nd.manual_file_path,
                      a.candidate_id, c.full_name AS candidate_name, c.email AS candidate_email,
                      r.title AS job_title, r.id AS requisition_id
               FROM nda_documents nd
               JOIN applications a ON a.id = nd.application_id
               JOIN candidates c ON c.id = a.candidate_id
               JOIN requisitions r ON r.id = a.requisition_id
               WHERE ($1::text IS NULL OR nd.status = $1)
               ORDER BY nd.created_at DESC""",
            status,
        )
    return [dict(r) for r in rows]


@nda_router.get("/{nda_id}/manual-file")
async def download_manual_nda_file(nda_id: str, actor: Actor = Depends(get_actor)):
    async with db.tenant_conn(actor.tenant_id) as conn:
        row = await conn.fetchrow("SELECT manual_file_path FROM nda_documents WHERE id=$1", nda_id)
    if not row or not row["manual_file_path"]:
        raise HTTPException(404, "No uploaded file for this NDA")
    abs_path = Path("/app") / row["manual_file_path"].lstrip("/")
    if not abs_path.exists():
        raise HTTPException(404, "File missing from disk")
    return FileResponse(str(abs_path), filename=abs_path.name)


# ─── Auto-advance + notify KAE / team lead ──────────────────────────────────

async def _on_nda_signed(tenant_id: str, application_id: str):
    """Advance stage -> screened and notify the client's KAE(s) + the
    assigned recruiter's team lead so they can schedule internal screening.
    """
    async with db.tenant_conn(tenant_id) as conn:
        old = await conn.fetchrow(
            "SELECT stage, requisition_id, assigned_recruiter_id FROM applications WHERE id=$1", application_id)
        if old is None or old["stage"] == "screened":
            return

        row = await conn.fetchrow(
            "UPDATE applications SET stage='screened', updated_at=now() WHERE id=$1 RETURNING updated_at",
            application_id,
        )
        await events.write_outbox(
            conn, tenant_id, "application.stage_changed",
            {"application_id": application_id, "from": old["stage"], "to": "screened", "reason": "nda_signed"},
            f"application.stage_changed:{application_id}:{row['updated_at'].isoformat()}",
        )

        req = await conn.fetchrow("SELECT client_id, title FROM requisitions WHERE id=$1", old["requisition_id"])
        cand = await conn.fetchrow(
            "SELECT c.full_name FROM applications a JOIN candidates c ON c.id=a.candidate_id WHERE a.id=$1",
            application_id)
        cand_name = cand["full_name"] if cand else "Candidate"
        job_title = req["title"] if req else "a role"

        recipient_ids: set[str] = set()
        if req and req["client_id"]:
            kae_rows = await conn.fetch(
                "SELECT user_id FROM client_owners WHERE tenant_id=$1 AND client_id=$2 "
                "AND owner_type='kae' AND is_active", tenant_id, req["client_id"])
            recipient_ids.update(str(r["user_id"]) for r in kae_rows)

        team_lead_found = False
        if old["assigned_recruiter_id"]:
            mgr = await conn.fetchval("SELECT reporting_to FROM users WHERE id=$1", old["assigned_recruiter_id"])
            if mgr:
                recipient_ids.add(str(mgr))
                team_lead_found = True
        if not team_lead_found:
            mgr_rows = await conn.fetch(
                "SELECT id FROM users WHERE tenant_id=$1 AND role='manager' AND is_active", tenant_id)
            recipient_ids.update(str(r["id"]) for r in mgr_rows)

        title = f"NDA signed - {cand_name} ready for internal screening"
        body_txt = (
            f"{cand_name} has signed their NDA for {job_title} and moved to Screened. "
            f"Please schedule an internal screening video call before submitting to the client."
        )
        for uid in recipient_ids:
            await conn.execute(
                """INSERT INTO notifications
                     (tenant_id, user_id, recipient_user_id, title, body, type, resource, resource_id, channel)
                   VALUES ($1,$2,$2,$3,$4,'info','application',$5,'inapp')""",
                tenant_id, uid, title, body_txt, application_id,
            )

    if recipient_ids:
        import asyncio
        asyncio.create_task(_email_recipients(tenant_id, recipient_ids, title, body_txt))


async def _email_recipients(tenant_id: str, user_ids: set[str], subject: str, body_text: str):
    # users has FORCE RLS — must go through db.tenant_conn (sets app.tenant_id),
    # not a raw asyncpg.connect() like the email_settings lookups above (that
    # table has no RLS, so an unscoped connection is fine there but not here).
    try:
        async with db.tenant_conn(tenant_id) as conn:
            rows = await conn.fetch(
                "SELECT email, full_name FROM users WHERE id = ANY($1::uuid[])", list(user_ids))
    except Exception as exc:
        print(f"NDA notify: could not resolve recipient emails: {exc}")
        return
    for r in rows:
        if r["email"]:
            await _send_email_with_pdf(tenant_id, r["email"], r["full_name"], subject, body_text)


# ─── Public signing endpoints (no auth — SECURITY DEFINER bypasses RLS) ─────

@nda_sign_public.get("/public")
async def get_nda_for_signing(token: str):
    async with db.system_conn() as conn:
        row = await conn.fetchrow("SELECT * FROM get_nda_by_signing_token($1)", token)
    if not row:
        raise HTTPException(404, "Signing link is invalid or expired")
    if row["status"] == "e_signed":
        return {"already_signed": True, "candidate_name": row["candidate_name"], "company_name": row["company_name"]}
    return {
        "already_signed": False,
        "candidate_name": row["candidate_name"],
        "job_title": row["job_title"],
        "company_name": row["company_name"],
        "letter_text": row["final_text"] or row["draft_text"],
        "otp_required": row["sign_method"] == "otp",
        "has_attached_file": row["attachment_source"] != "generated",
        "attached_file_name": row["attached_file_name"],
    }


@nda_sign_public.get("/attached-file")
async def download_attached_file(token: str):
    """Public, token-gated download of the custom NDA/Contract file (if the
    recruiter chose to attach one instead of the auto-generated PDF)."""
    async with db.system_conn() as conn:
        row = await conn.fetchrow("SELECT * FROM get_nda_attached_file_by_token($1)", token)
    if not row or not row["attached_file_path"]:
        raise HTTPException(404, "No attached file for this signing link")
    abs_path = Path("/app") / row["attached_file_path"].lstrip("/")
    if not abs_path.exists():
        raise HTTPException(404, "File missing from disk")
    return FileResponse(str(abs_path), filename=row["attached_file_name"] or abs_path.name)


@nda_sign_public.post("/request-otp")
async def request_nda_otp(token: str):
    async with db.system_conn() as conn:
        row = await conn.fetchrow("SELECT * FROM request_nda_otp_by_token($1)", token)
    if not row:
        raise HTTPException(400, "Unable to send OTP for this signing link")
    import asyncio
    asyncio.create_task(_send_email_with_pdf(
        str(row["tenant_id"]), row["candidate_email"], row["candidate_name"],
        "Your NDA signing verification code",
        f'Dear {row["candidate_name"]},\n\nYour verification code is: {row["otp_code"]}\n\n'
        f'This code expires in 10 minutes.',
    ))
    return {"sent": True}


class NdaSignRequest(BaseModel):
    signatory_name: str
    otp_code: Optional[str] = None


@nda_sign_public.post("/sign")
async def sign_nda(token: str, body: NdaSignRequest):
    name = (body.signatory_name or "").strip()
    if not name:
        raise HTTPException(400, "Please enter your full name as a signature")

    async with db.system_conn() as conn:
        info = await conn.fetchrow("SELECT * FROM get_nda_by_signing_token($1)", token)
        if not info:
            raise HTTPException(404, "Signing link is invalid or expired")
        if info["sign_method"] == "otp":
            if not body.otp_code:
                raise HTTPException(400, "OTP code is required")
            ok = await conn.fetchval("SELECT verify_nda_otp_by_token($1, $2)", token, body.otp_code)
            if not ok:
                raise HTTPException(400, "Invalid or expired OTP code")
        result = await conn.fetchrow("SELECT * FROM sign_nda_by_token($1, $2)", token, name)
    if not result:
        raise HTTPException(400, "Signing link is invalid, already used, or expired")

    await _on_nda_signed(str(result["tenant_id"]), str(result["application_id"]))
    return {"signed": True, "message": "Thank you! Your e-signature has been recorded."}
